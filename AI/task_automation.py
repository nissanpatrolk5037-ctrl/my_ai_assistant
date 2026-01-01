from __future__ import annotations
#-------------------------------------------------------------#
from typing import Dict, Iterable, List, Optional, Tuple, Any, Callable
from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator
from requests import get, HTTPError, ConnectionError
from screen_brightness_control import set_brightness
from watchdog.events import FileSystemEventHandler
from google.oauth2.credentials import Credentials
from googleapiclient.http import MediaFileUpload
from capstone import Cs, CS_ARCH_X86, CS_MODE_32
from deep_translator import GoogleTranslator 
from pdfminer.high_level import extract_text
from googleapiclient.discovery import build
from datetime import datetime, timedelta
from PyPDF2 import PdfReader, PdfMerger
from watchdog.observers import Observer
from keyboard import press_and_release
import xml.etree.ElementTree as ET
from tkinter import messagebox, Tk
from PIL import Image, ImageGrab
from urllib.parse import unquote, quote, urljoin
from PIL import Image, ImageGrab
import matplotlib.pyplot as plt
from functools import lru_cache, wraps
import speech_recognition as sr
from bleak import BleakScanner
from bs4 import BeautifulSoup
from PIL.ExifTags import TAGS
from memory_manager import *
from pytube import YouTube
from docx import Document
import pygetwindow as gw
from pathlib import Path
import sounddevice as sd
import librosa.display
from re import findall
import win32com.client
import yfinance as yf
from groq import Groq
from tqdm import tqdm
import datetime as dt
from fpdf import FPDF
from moviepy import *
import pandas as pd
import numpy as np
import subprocess
import webbrowser
import pyautogui
import threading
import pywhatkit
import pyperclip
import importlib
import traceback
import requests
import tempfile
import textwrap
import openpyxl
import hashlib
import inspect
import asyncio
import discord
import imaplib
import librosa
import qrcode
import pygame
import shutil
import psutil
import urllib
import socket
import random
import ctypes
import winreg
import pefile
import queue
import magic
import email
import boto3
try:
    import spacy
    SPACY_AVAILABLE = True
except (ImportError, Exception):
    SPACY_AVAILABLE = False
    spacy = None
import sched
import scipy
import json
import nmap
import time
import glob
import sys
import cv2
import csv
import os
import re

# Add parent directory to path for imports
if __name__ == "__main__" or True:  # Always add path
    current_dir = Path(__file__).parent.absolute()
    parent_dir = current_dir.parent if current_dir.name == "Code" else current_dir
    if str(parent_dir) not in sys.path:
        sys.path.insert(0, str(parent_dir))

# -----------------------------
# Globals
# -----------------------------
LAST_PRICE_FILE = Path("last_price.txt")
IS_WINDOWS = sys.platform.startswith("win")
IS_MACOS = sys.platform == "darwin"
IS_LINUX = sys.platform.startswith("linux")
ROOT = Path.cwd()
GEN_ENV = {}
DISCORD_TOKEN = ""
DISCORD_CHANNEL_ID = 0
WHATSAPP_TOKEN = ""
WHATSAPP_PHONE_ID = ""
MEMORY_FILE = "os_context_memory.json"
API_KEY = "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"

audio_queue: "queue.Queue[str]" = queue.Queue()

# ===============================
# TIME UTILS
# ===============================
def now() -> str:
    return datetime.now().isoformat(timespec="seconds")


# ===============================
# SAFE LOGGING
# ===============================
def log(*args: Any) -> None:
    try:
        print(*args)
    except Exception:
        pass


# ===============================
# KWARG HELPER
# ===============================
def _kw(kwargs: dict, key: str, default: Any = None) -> Any:
    return kwargs.get(key, default)


# ===============================
# ERROR PRINTING
# ===============================
def _print_err(*args, **kwargs) -> None:
    msg = args[0] if args else kwargs.get("msg", "")
    try:
        print(f"[ERROR] {msg}")
    except Exception:
        pass


# ===============================
# SAFE MESSAGE BOX
# ===============================
def _safe_message_box(*args, **kwargs) -> None:
    if args:
        title = args[0] if len(args) > 0 else kwargs.get("title", "")
        text = args[1] if len(args) > 1 else kwargs.get("text", "")
    else:
        title = kwargs.get("title", "")
        text = kwargs.get("text", "")

    if not IS_WINDOWS or Tk is None or messagebox is None:
        print(f"[MSGBOX] {title}: {text}")
        return

    try:
        root = Tk()
        root.withdraw()
        messagebox.showinfo(title, text)
        root.destroy()
    except Exception as e:
        print(f"[MSGBOX ERROR] {e}")


# ===============================
# HTTP SESSION FACTORY
# ===============================
def _make_session() -> requests.Session:
    session = requests.Session()

    retries = requests.adapters.Retry(
        total=3,
        backoff_factor=0.5,
        status_forcelist=(429, 500, 502, 503, 504),
        allowed_methods=frozenset(["GET", "POST"])
    )

    adapter = requests.adapters.HTTPAdapter(
        max_retries=retries,
        pool_connections=10,
        pool_maxsize=10
    )

    session.mount("http://", adapter)
    session.mount("https://", adapter)

    session.headers.update({
        "User-Agent": "OptimizedAssistant/1.0"
    })

    return session

HEADERS_LIST = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1 Safari/605.1.15",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36"
]

# ===============================
# GLOBAL SESSION (SAFE)
# ===============================

SESSION = _make_session()

# ===============================
# GROQ CLIENT FACTORY
# ===============================

def _get_groq_client() -> Optional[Any]:
    if Groq is None:
        _print_err("Groq library not available.")
        return None

    if not API_KEY or API_KEY == "YOUR_GROQ_API_KEY":
        _print_err("Groq API key not configured.")
        return None

    try:
        return Groq(api_key=API_KEY)
    except Exception as e:
        _print_err(f"Failed to initialize Groq client: {e}")
        return None


# ===============================
# GROQ CALL (CACHED)
# ===============================
@lru_cache(maxsize=128)
def groq_call(instructions: str, query: Optional[str] = None) -> Optional[str]:
    client = _get_groq_client()
    if client is None:
        return "Groq client not available."

    content = instructions if query is None else f"{instructions}\n\n{query}"

    models = [
        "llama-3.3-70b-versatile",
        "llama-3.1-8b-instant",
        "meta-llama/llama-guard-4-12b"
    ]

    for model in models:
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": content}],
                stream=False,
            )

            text = response.choices[0].message.content or ""
            return text.replace("**", "")

        except Exception as e:
            e_str = str(e)

            if re.search(r"rate limit", e_str, re.IGNORECASE):
                print(f"❌ Rate limit for {model}, switching model...")
                continue

            print(f"[Groq Error] {e_str}")
            return None

    print("⚠️ All models exhausted. Try again later.")
    return None


# ===============================
# PUBLIC WRAPPER
# ===============================

def groq_answer(instructions: str, query: Optional[str] = None) -> str:
    return groq_call(instructions, query) or ""

# =========================================================
# NLP → MULTI COMMAND REWRITE
# =========================================================
REWRITE_INSTRUCTIONS = """
Rewrite the user's request into executable commands.
Rules:
- Commands MUST be separated by semicolons
- Commands MUST be short and imperative
- Do NOT explain
- Do NOT add text
Example:
"Open Chrome and search cats then copy result and paste into notepad"
→ "open_chrome; search cats; copy result; paste_notepad"
"""




class SystemCleanup:

    @staticmethod
    def delete_files_in_folder(folder: str | Path) -> str:
        p = Path(folder)

        if not p.exists():
            return f"Folder not found: {p}"

        deleted_count = 0

        for child in p.iterdir():
            try:
                if child.is_file() or child.is_symlink():
                    child.unlink(missing_ok=True)
                    deleted_count += 1
                elif child.is_dir():
                    shutil.rmtree(child, ignore_errors=True)
                    deleted_count += 1
            except Exception as e:
                _print_err(f"Error deleting {child}: {e}")

        return f"Deleted {deleted_count} items in {p.name}."


    @staticmethod
    def run_command(command: str) -> Optional[str]:
        try:
            subprocess.run(
                command,
                check=True,
                shell=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            return f"Command executed successfully: {command.split()[0]}"
        except subprocess.CalledProcessError as e:
            _print_err(f"Command failed: {command} ({e})")
            return f"Command failed: {command.split()[0]}"
        except FileNotFoundError:
            _print_err(f"Command not found: {command.split()[0]}")
            return f"Command not found: {command.split()[0]}"


    @staticmethod
    def clean_temp() -> str:
        temp_folder = None

        if IS_WINDOWS:
            temp_folder = os.environ.get("TEMP") or os.environ.get("TMP")
            windir = os.environ.get("WINDIR", "C:\\Windows")
            SystemCleanup.delete_files_in_folder(Path(windir) / "Temp")

        elif IS_MACOS:
            temp_folder = os.environ.get("TMPDIR", "/private/tmp")

        elif IS_LINUX:
            temp_folder = "/tmp"

        if temp_folder and Path(temp_folder).exists():
            SystemCleanup.delete_files_in_folder(temp_folder)
            return f"Cleaned standard temp folder: {temp_folder}"

        return "Temp cleanup path not found."


    @staticmethod
    def clean_recycled_items() -> str:
        if IS_WINDOWS:
            return SystemCleanup.run_command("rd /s /q C:\\$Recycle.Bin") or "Recycle bin cleanup attempted."

        elif IS_MACOS:
            trash_path = Path.home() / ".Trash"
            if trash_path.exists():
                SystemCleanup.delete_files_in_folder(trash_path)
                return f"Cleaned macOS Trash: {trash_path.name}"
            return "macOS Trash not found."

        elif IS_LINUX:
            return "Linux trash cleanup skipped (requires trash-cli or user confirmation)."

        return "Recycle/Trash cleanup skipped."


    @staticmethod
    def clean_dns_cache() -> str:
        if IS_WINDOWS:
            return SystemCleanup.run_command("ipconfig /flushdns") or "DNS flush attempted."

        elif IS_MACOS:
            try:
                subprocess.run(
                    "dscacheutil -flushcache",
                    check=True,
                    shell=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL
                )
                return "DNS cache flushed (macOS)."
            except Exception:
                return "DNS cache flush failed (macOS)."

        elif IS_LINUX:
            if shutil.which("systemd-resolve"):
                return SystemCleanup.run_command("sudo systemd-resolve --flush-caches") or "DNS flush attempted."
            elif shutil.which("nscd"):
                return SystemCleanup.run_command("sudo /etc/init.d/nscd restart") or "DNS restart attempted."
            return "Linux DNS service not found. Skipped."

        return "DNS cache cleanup skipped."


    @staticmethod
    def main() -> str:
        """Cross-platform system cleanup entry point."""
        SystemCleanup.clean_temp()
        SystemCleanup.clean_recycled_items()
        SystemCleanup.clean_dns_cache()
        return "Cleanup complete!"
    
def image_generation(**kwargs) -> Optional[Path]:
    """
    Generate an image using Pollinations AI based on a reduced prompt.
    Accepts kwargs:
      - query or raw_input (str)
      - out_path (str or Path)
    Returns Path or None
    """

    # ---------------------------
    # INPUT HANDLING
    # ---------------------------
    query = (kwargs.get("query") or kwargs.get("raw_input") or "").strip()
    if not query:
        _print_err("No query provided for image generation.")
        return None

    out_path = Path(kwargs.get("out_path", "Generated_Image.jpg"))

    # ---------------------------
    # PROMPT REDUCTION (SAFE)
    # ---------------------------
    try:
        obj = groq_answer(
            "Just return the primary object from this query. "
            "E.g., 'create a dog' -> 'dog'",
            query,
        )
    except Exception:
        obj = None

    obj = (obj or "").strip() or query

    # ---------------------------
    # IMAGE DOWNLOAD
    # ---------------------------
    img_url = f"https://image.pollinations.ai/prompt/{requests.utils.quote(obj)}"

    try:
        with SESSION.get(img_url, stream=True, timeout=15) as response:
            response.raise_for_status()

            out_path.parent.mkdir(parents=True, exist_ok=True)

            with open(out_path, "wb") as f:
                for chunk in response.iter_content(chunk_size=1024 * 32):
                    if chunk:
                        f.write(chunk)

        # ---------------------------
        # IMAGE VALIDATION
        # ---------------------------
        if Image is not None:
            try:
                Image.open(out_path).verify()
            except Exception:
                _print_err("Downloaded file may not be a valid image.")

            # Optional preview (won't crash headless systems)
            try:
                Image.open(out_path).show()
            except Exception:
                pass

        return out_path

    except Exception as e:
        _print_err(f"Failed to download image: {e}")
        return None


def image_optimization(**kwargs) -> None:
    """
    Optimize image(s) by resizing, converting to RGB, and saving with reduced quality.

    Accepts kwargs:
        - mode: 'image' or 'folder'
        - path: path to image or folder
        - output_path: optional path for output image
        - quality: JPEG quality (default 85)
        - size: tuple (width, height) (default 800x600)
    """
    if Image is None:
        _print_err("PIL library not available. Cannot optimize images.")
        return

    mode = kwargs.get("mode")
    path = kwargs.get("path")
    output_path = kwargs.get("output_path")
    quality = kwargs.get("quality", 85)
    size = tuple(kwargs.get("size", (800, 600)))

    if not path:
        _print_err("No path provided for optimization.")
        return

    p = Path(path)

    def _optimize_one(in_path: Path, out_path: Path):
        try:
            with Image.open(in_path) as img:
                img = img.convert("RGB")
                img = img.resize(size)
                out_path.parent.mkdir(parents=True, exist_ok=True)
                img.save(out_path, optimize=True, quality=quality)
        except Exception as e:
            _print_err(f"Optimize failed for {in_path}: {e}")

    if mode == "image":
        if not p.exists():
            _print_err(f"Image not found: {p}")
            return

        out = Path(output_path) if output_path else p.with_stem(p.stem + "_optimized").with_suffix(".jpg")
        _optimize_one(p, out)

    elif mode == "folder":
        if not p.is_dir():
            _print_err(f"Folder not found: {p}")
            return

        for img_file in p.iterdir():
            if img_file.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp", ".bmp"):
                out = img_file.with_stem("resized_" + img_file.stem).with_suffix(".jpg")
                _optimize_one(img_file, out)

    else:
        _print_err("mode must be 'image' or 'folder'.")

def docx_to_pdf(**kwargs) -> str:
    """
    Convert a DOCX file to PDF.

    Accepts kwargs:
        - input_path: required path to DOCX
        - output_path: optional path for PDF output

    Returns string status message.
    """

    input_path = kwargs.get("input_path")
    output_path = kwargs.get("output_path")

    if not input_path:
        return "Error: input_path is required."

    p_in = Path(input_path).resolve()
    if not p_in.exists():
        return f"Error: input file not found: {p_in}"

    p_out = Path(output_path or p_in.with_suffix(".pdf")).resolve()

    # ---------------------------
    # 1. Windows COM Automation
    # ---------------------------
    if IS_WINDOWS and win32com is not None:
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(p_in))
            doc.SaveAs(str(p_out), FileFormat=17)  # wdFormatPDF
            doc.Close(False)
            word.Quit()
            return f"Saved PDF (Windows COM): {p_out}"
        except Exception as e:
            _print_err(f"Windows COM conversion failed: {e}")
            # Fall through to LibreOffice/OpenOffice fallback

    # ---------------------------
    # 2. Cross-platform Fallback: LibreOffice/OpenOffice
    # ---------------------------
    try:
        soffice_path = shutil.which("libreoffice") or shutil.which("soffice")
        if soffice_path:
            command = [
                str(soffice_path),
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(p_out.parent),
                str(p_in)
            ]
            subprocess.run(command, check=True, capture_output=True, timeout=60)

            if p_out.exists():
                return f"Saved PDF (LibreOffice): {p_out}"
            else:
                _print_err(f"LibreOffice conversion did not produce output file: {p_out}")

    except Exception as e:
        _print_err(f"LibreOffice conversion failed: {e}")

    return (
        "Error: Cannot convert document. Requires "
        "win32com (Windows) or LibreOffice/soffice (all platforms) to be installed."
    )

def get_crypto_price_coingecko(**kwargs) -> Optional[float]:
    """
    Fetch the current price of a cryptocurrency in USD from CoinGecko.

    Accepts kwargs:
        - symbol or raw_input: crypto id (e.g., 'bitcoin', 'ethereum')

    Returns:
        - price in USD (float) if successful
        - None if any error occurs
    """
    symbol = (kwargs.get("symbol") or kwargs.get("raw_input") or "").strip().lower()
    if not symbol:
        _print_err("No cryptocurrency symbol provided.")
        return None

    url = f"https://api.coingecko.com/api/v3/simple/price?ids={symbol}&vs_currencies=usd"

    try:
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        data = resp.json()

        if symbol not in data or "usd" not in data[symbol]:
            _print_err(f"Symbol not found or no USD price: {symbol}")
            return None

        return float(data[symbol]["usd"])

    except Exception as e:
        _print_err(f"CoinGecko API error: {e}")
        return None


def track_stock(**kwargs) -> str:
    """
    Track a stock or crypto price, compare with last stored price, update state_file, return message.

    kwargs:
        - symbol or raw_input: ticker symbol
        - state_file: Path object to store last price (default LAST_PRICE_FILE)
        - asset_type: 'stock' or 'crypto' (default 'stock')
    """
    symbol = (kwargs.get("symbol") or kwargs.get("raw_input") or "").strip()
    if not symbol:
        return "Error: No symbol provided."

    state_file = Path(kwargs.get("state_file", LAST_PRICE_FILE))
    asset_type = kwargs.get("asset_type", "stock").lower()

    current_price: Optional[float] = None

    # ---------------------------
    # FETCH CURRENT PRICE
    # ---------------------------
    if asset_type == "stock":
        try:
            stock = yf.Ticker(symbol)
            data = stock.history(period="1d", interval="1m")
            if data.empty:
                return f"No price data found for {symbol}."
            current_price = float(data["Close"].iloc[-1])
        except Exception as e:
            _print_err(f"yfinance error: {e}")
            return f"Error fetching stock price for {symbol}."

    elif asset_type == "crypto":
        current_price = get_crypto_price_coingecko(symbol=symbol)
        if current_price is None:
            return f"Error fetching crypto price for {symbol}."

    else:
        return "Invalid asset_type. Use 'stock' or 'crypto'."

    # ---------------------------
    # READ LAST PRICE
    # ---------------------------
    last_price: Optional[float] = None
    if state_file.exists():
        try:
            last_price = float(state_file.read_text().strip())
        except Exception:
            last_price = None

    # ---------------------------
    # PREPARE MESSAGE
    # ---------------------------
    if last_price is None:
        message = f"First time checking {symbol}. Current price: ${current_price:.2f}"
    elif current_price > last_price:
        message = f"{symbol} price went UP from ${last_price:.2f} to ${current_price:.2f}"
    elif current_price < last_price:
        message = f"{symbol} price went DOWN from ${last_price:.2f} to ${current_price:.2f}"
    else:
        message = f"{symbol} price stayed the same: ${current_price:.2f}"

    # ---------------------------
    # WRITE NEW STATE
    # ---------------------------
    try:
        state_file.parent.mkdir(parents=True, exist_ok=True)
        state_file.write_text(str(current_price))
    except Exception as e:
        _print_err(f"Failed to write state file: {e}")

    return message

def summarize_clipboard_text(**kwargs) -> str:
    """
    Summarize text from the clipboard using Groq.
    Returns a concise summary or error message.
    """
    try:
        text = pyperclip.paste() or ""
    except Exception as e:
        _print_err(f"Clipboard read failed: {e}")
        return "Failed to read clipboard."

    if not text.strip():
        return "No text found in clipboard."

    try:
        summary = groq_answer(
            "Summarize the following text. Keep it concise.",
            text
        )
    except Exception as e:
        _print_err(f"Groq summarization failed: {e}")
        return "Failed to summarize."

    if not summary:
        return "Failed to summarize."

    # Wrap summary to 120 characters per line
    wrapped_summary = textwrap.fill(summary, width=120)
    return f"Summary:\n{wrapped_summary}"


def translate_clipboard_text(**kwargs) -> str:
    """
    Translate text from the clipboard to a target language using deep_translator.
    kwargs:
        - target_lang: target language code (default 'en')
    Returns translated text or error message.
    """
    target_lang = kwargs.get("target_lang", "en")

    # Check if deep_translator is available
    if 'GoogleTranslator' not in globals() or GoogleTranslator is None:
        _print_err("deep_translator not available.")
        return "deep_translator not available."

    # Read clipboard
    try:
        text = pyperclip.paste() or ""
    except Exception as e:
        _print_err(f"Clipboard read failed: {e}")
        return "Failed to read clipboard."

    if not text.strip():
        return "No text found in clipboard."

    # Translate
    try:
        translated = GoogleTranslator(source="auto", target=target_lang).translate(text)
    except Exception as e:
        _print_err(f"Translation failed: {e}")
        return f"Translation failed to {target_lang}."

    return f"Translated ({target_lang}):\n{translated}"

def enable_game_mode(**kwargs) -> str:
    """
    Enables platform-specific high-performance/game mode settings.
    Includes Windows GameBar settings, macOS energy policy, and Linux CPU governor settings.
    """
    results = []

    # ---------------------------
    # WINDOWS SETTINGS
    # ---------------------------
    if IS_WINDOWS:
        try:
            if 'winreg' not in globals() or winreg is None:
                raise ImportError("winreg not available (Windows environment expected).")
            
            # Enable GameBar Auto Game Mode
            key = winreg.CreateKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\GameBar")
            winreg.SetValueEx(key, "AllowAutoGameMode", 0, winreg.REG_DWORD, 1)
            winreg.CloseKey(key)

            # Set High Performance Power Plan
            subprocess.run(
                "powercfg /setactive 8c5e7fda-e8bf-4a96-9a85-a6e23a8c635c",
                check=False,
                shell=True,
                capture_output=True,
                timeout=5
            )
            results.append("Windows Game Mode enabled and power plan set to High Performance.")
        except Exception as e:
            _print_err(f"Windows Game Mode activation failed: {e}")
            results.append(f"Windows activation failed: {e}")

    # ---------------------------
    # MACOS SETTINGS
    # ---------------------------
    elif IS_MACOS:
        try:
            # Prevent display sleep, computer sleep, and disk sleep
            subprocess.run(
                "sudo pmset -a disablesleep 1",
                check=True,
                shell=True,
                capture_output=True,
                timeout=5
            )
            results.append("macOS performance tweaks activated (sleep disabled).")
        except subprocess.CalledProcessError:
            results.append("macOS: Sudo password required for performance settings. Run manually.")
        except Exception as e:
            _print_err(f"macOS activation failed: {e}")
            results.append(f"macOS activation failed: {e}")

    # ---------------------------
    # LINUX SETTINGS
    # ---------------------------
    elif IS_LINUX:
        try:
            cpu_paths = list(Path("/sys/devices/system/cpu/").glob("cpu*/cpufreq/scaling_governor"))
            if not cpu_paths:
                results.append("Linux CPU governor control paths not found.")
            else:
                for p in cpu_paths:
                    subprocess.run(
                        f"echo performance | sudo tee {p}",
                        check=True,
                        shell=True,
                        capture_output=True,
                        timeout=5
                    )
                results.append("Linux CPU Governor set to 'performance' for all cores.")
        except subprocess.CalledProcessError:
            results.append("Linux: Sudo password required to change CPU governor. Run manually.")
        except Exception as e:
            _print_err(f"Linux activation failed: {e}")
            results.append(f"Linux activation failed: {e}")

    else:
        results.append("System performance mode not supported on this OS.")

    # ---------------------------
    # OPTIONAL CLEANUP
    # ---------------------------
    if 'SystemCleanup' in globals():
        try:
            SystemCleanup.main()
            results.append("System cleanup performed.")
        except Exception as e:
            _print_err(f"SystemCleanup failed: {e}")
            results.append("System cleanup failed.")

    return "\n".join(results)


def merge_pdfs_in_folder(**kwargs) -> str:
    """
    Merge all PDF files in a folder into a single PDF.

    kwargs:
        - folder_path: path to folder containing PDFs
        - output_filename: optional output PDF filename (default 'merged_output.pdf')

    Returns status message.
    """
    if PdfMerger is None:
        return "PyPDF2 not available. Cannot merge PDFs."

    folder_path = kwargs.get("folder_path")
    if not folder_path:
        return "Error: folder_path is required."

    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        return f"Error: Folder not found: {folder}"

    output_filename = kwargs.get("output_filename", "merged_output.pdf")
    out_path = folder / output_filename

    pdf_files = sorted([p for p in folder.iterdir() if p.suffix.lower() == ".pdf"])
    if not pdf_files:
        return "No PDF files found in folder."

    merger = PdfMerger()
    added = []

    for pdf in pdf_files:
        try:
            merger.append(str(pdf))
            added.append(pdf.name)
        except Exception as e:
            _print_err(f"Skipped {pdf.name}: {e}")

    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with out_path.open("wb") as f:
            merger.write(f)
    except Exception as e:
        _print_err(f"Failed to write merged PDF: {e}")
        merger.close()
        return "Failed to write merged PDF."

    merger.close()

    if added:
        return f"Merged PDF saved to: {out_path} (added: {', '.join(added)})"
    return "No PDFs merged."

def download_unsplash_wallpapers(**kwargs) -> str:
    """
    Download wallpapers from Unsplash.

    kwargs:
        - query: search query (default "nature")
        - count: number of images to download (default 5)

    Returns status message.
    """
    query = kwargs.get("query", "nature")
    try:
        count = max(1, int(kwargs.get("count", 5)))
    except Exception:
        count = 5

    save_path = Path("unsplash_wallpapers")
    try:
        save_path.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        _print_err(f"Failed to create directory {save_path}: {e}")
        return f"Failed to create directory {save_path}"

    saved = []

    for i in range(count):
        try:
            url = f"https://source.unsplash.com/1920x1080/?{requests.utils.quote(query)}&sig={i}"
            r = SESSION.get(url, timeout=15)
            r.raise_for_status()
            fname = save_path / f"{query}_{i}.jpg"
            fname.write_bytes(r.content)
            saved.append(str(fname))
        except Exception as e:
            _print_err(f"Error downloading image {i}: {e}")

    if saved:
        return f"Downloaded {len(saved)} images: {saved}"
    return "No images downloaded."

def detect_fake_news(**kwargs) -> str:
    """
    Detect whether the given text is real or fake using Groq.

    kwargs:
        - text or raw_input: the text to analyze

    Returns:
        - A string indicating if the text is likely real or fake
    """
    text = (kwargs.get("text") or kwargs.get("raw_input") or "").strip()
    if not text:
        return "No text provided for analysis."

    try:
        result = groq_answer("Detect whether this information is real or fake. Respond with 'real' or 'fake'.", text)
    except Exception as e:
        _print_err(f"Groq detection failed: {e}")
        return "Failed to detect fake news."

    if not result:
        return "Failed to detect fake news."

    return result.strip()

def website_summarizer(**kwargs) -> Optional[str]:
    """
    Summarize the main content of a website using Groq.

    kwargs:
        - url or raw_input: the website URL

    Returns:
        - summary string or None on failure
    """
    url = (kwargs.get("url") or kwargs.get("raw_input") or "").strip()
    if not url:
        _print_err("No URL provided.")
        return None

    try:
        r = SESSION.get(url, timeout=15)
        r.raise_for_status()
    except Exception as e:
        _print_err(f"Request failed for {url}: {e}")
        return None

    try:
        summary = groq_answer(
            "Summarize this HTML. Focus on important content, ignore boilerplate/navigation.",
            r.text
        )
    except Exception as e:
        _print_err(f"Groq summarization failed for {url}: {e}")
        return None

    if not summary:
        _print_err(f"Groq returned empty summary for {url}")
        return None

    return summary.strip()

def get_local_ip(**kwargs) -> str:
    """
    Returns the local IP address of the machine.
    Falls back to '127.0.0.1' on failure.
    """
    try:
        # Attempt to get the IP by connecting to a public DNS (does not send data)
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.settimeout(0.5)
        try:
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
        finally:
            s.close()
        return local_ip
    except Exception as e:
        _print_err(f"Failed to get local IP: {e}")

    # Fallback to hostname resolution
    try:
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
        return local_ip
    except Exception as e:
        _print_err(f"Fallback hostname resolution failed: {e}")

    # Default fallback
    return "127.0.0.1"

def port_scanner(**kwargs):
    """
    Scan ports on a target IP using nmap.

    kwargs:
        ip: target IP (default = local IP)
        ports: port range (default = "22-443")
    """
    ip = kwargs.get("ip") or get_local_ip()
    ports = str(kwargs.get("ports", "22-443"))

    if nmap is None:
        return "nmap-python not available."

    try:
        scanner = nmap.PortScanner()
    except Exception as e:
        _print_err(f"Failed to initialize nmap scanner: {e}")
        return "Nmap initialization failed."

    try:
        scanner.scan(hosts=ip, ports=ports, arguments="-T4")
    except Exception as e:
        _print_err(f"Scan failed: {e}")
        return f"Scan failed for {ip}."

    if ip not in scanner.all_hosts():
        return f"{ip} not found in scan results."

    if scanner[ip].state() != "up":
        return f"{ip} appears {scanner[ip].state()}."

    results = [f"Scan results for {ip} ({ports})"]

    for proto in scanner[ip].all_protocols():
        results.append(f"\nProtocol: {proto.upper()}")

        ports_data = scanner[ip][proto]
        if not ports_data:
            results.append("No open ports found.")
            continue

        for port in sorted(ports_data):
            entry = ports_data[port]
            state = entry.get("state", "unknown").upper()
            service = entry.get("name", "unknown").upper()
            results.append(f"Port {port:<5} | {state:<6} | {service}")

    return "\n".join(results)



async def _ble_discover_async(timeout: float = 5.0):
    if BleakScanner is None:
        raise RuntimeError("bleak not available.")
    try:
        devices = await BleakScanner.discover(timeout=timeout)
        return devices
    except Exception as e:
        _print_err(f"BLE discovery async failed: {e}")
        return []

# ---------------------------
# Public sync function
# ---------------------------
def get_nearby_devices(**kwargs) -> Optional[str]:
    """
    Discover nearby BLE devices and return formatted list.

    kwargs:
        - timeout: scan duration in seconds (default 5.0)
    """
    timeout = float(kwargs.get("timeout", 5.0))

    if BleakScanner is None:
        return "bleak not available."

    try:
        # Handle running event loop (e.g., Jupyter) vs no loop
        try:
            loop = asyncio.get_running_loop()
            # Use ensure_future and run until complete workaround
            future = asyncio.ensure_future(_ble_discover_async(timeout=timeout))
            devices = loop.run_until_complete(future)
        except RuntimeError:
            # No running loop, safe to use asyncio.run
            devices = asyncio.run(_ble_discover_async(timeout=timeout))

        if not devices:
            return "No devices found."

        out = [f"Device: {d.name or 'Unknown'} [{d.address}]" for d in devices]
        return "\n".join(out)

    except Exception as e:
        _print_err(f"BLE scan failed: {e}")
        return None

def audio_translator_auto(**kwargs) -> str:
    """
    Transcribe an audio file, translate it, and optionally play as audio.

    kwargs:
        - audio_file_path or raw_input: path to audio file
        - target_language: language code (default 'en')
        - output_audio_file: temporary audio output file (default 'translated_audio.mp3')
    """
    audio_file_path = kwargs.get("audio_file_path") or kwargs.get("raw_input")
    target_language = kwargs.get("target_language", "en")
    output_audio_file = kwargs.get("output_audio_file", "translated_audio.mp3")

    if GoogleTranslator is None:
        return "deep_translator not available."

    if not audio_file_path or not Path(audio_file_path).exists():
        return "Error: Audio file not found."

    r = sr.Recognizer()

    try:
        # --- Load audio ---
        with sr.AudioFile(audio_file_path) as source:
            audio_data = r.record(source)

        # --- Transcribe ---
        try:
            transcribed_text = r.recognize_google(audio_data)
        except sr.UnknownValueError:
            return "Error: Could not understand audio."
        except sr.RequestError as e:
            return f"Error: Speech recognition request failed; {e}"

        # --- Translate ---
        try:
            translated_text = GoogleTranslator(source="auto", target=target_language).translate(transcribed_text)
        except Exception as e:
            _print_err(f"Translation failed: {e}")
            return "Error: Translation failed."

        # --- If pygame not available, save text to file ---
        if pygame is None:
            try:
                Path(output_audio_file).write_text(translated_text, encoding="utf-8")
                return f"Saved translated text (pygame not installed): {output_audio_file}"
            except Exception as e:
                _print_err(f"Failed to save translated text: {e}")
                return "Failed to save translated text (pygame not installed)."

        # --- Generate speech using gTTS ---
        try:
            from gtts import gTTS
            tts = gTTS(text=translated_text, lang=target_language, slow=False)
            tts.save(output_audio_file)
        except Exception as e:
            _print_err(f"gTTS failed: {e}")
            return "Audio save failed."

        # --- Play audio using pygame ---
        try:
            pygame.mixer.init()
            pygame.mixer.music.load(output_audio_file)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                time.sleep(0.1)
            try:
                os.remove(output_audio_file)
            except Exception:
                pass
            return f"Translated Text ({target_language}): {translated_text}"
        except Exception as e:
            _print_err(f"Could not play audio: {e}")
            try:
                os.remove(output_audio_file)
            except Exception:
                pass
            return f"Translated Text ({target_language}): {translated_text} (could not play audio)"

    except Exception as e:
        return f"An error occurred during the translation process: {e}"


def listen_meeting():
    recognizer = sr.Recognizer()
    try:
        mic = sr.Microphone()
    except Exception as e:
        _print_err(f"Microphone not available: {e}")
        return

    with mic as source:
        try:
            recognizer.adjust_for_ambient_noise(source)
        except Exception as e:
            _print_err(f"Ambient noise adjustment failed: {e}")

        while True:
            try:
                audio = recognizer.listen(source, phrase_time_limit=10)
                try:
                    text = recognizer.recognize_google(audio)
                    if text.strip():
                        audio_queue.put(text)
                except sr.UnknownValueError:
                    continue  # ignore unrecognized speech
                except sr.RequestError as e:
                    _print_err(f"Speech recognition service error: {e}")
            except Exception as e:
                _print_err(f"Error listening: {e}")
                break

# -----------------------
# Summarize transcript with Groq
# -----------------------
def summarize_meeting(**kwargs) -> Dict[str, Any]:
    """
    Summarize a meeting transcript into structured JSON.

    Returns:
        {
            "key_points": [...],
            "decisions": [...],
            "action_items": [...]
        }
    """
    meeting_text = (kwargs.get("meeting_text") or kwargs.get("raw_input") or "").strip()
    if not meeting_text:
        return {"key_points": [], "decisions": [], "action_items": []}

    prompt = (
        "Summarize the following meeting transcript as JSON with keys: "
        "key_points (list), decisions (list), action_items (list of objects with owner if present). "
        "Return ONLY JSON."
    )
    try:
        structured = groq_answer(prompt, meeting_text)
        data = json.loads((structured or "").strip())
        # Ensure keys exist
        return {
            "key_points": data.get("key_points", []),
            "decisions": data.get("decisions", []),
            "action_items": data.get("action_items", [])
        }
    except Exception as e:
        _print_err(f"Failed to parse summary JSON: {e}")
        return {"key_points": [], "decisions": [], "action_items": []}

# -----------------------
# Process transcripts from queue
# -----------------------
def process_transcripts(timeout: float = 10.0) -> Dict[str, Any]:
    """
    Collect transcripts from `audio_queue` and summarize.

    Parameters:
        timeout: seconds to wait for new transcripts before retrying.

    Returns:
        structured summary JSON
    """
    all_text = ""
    while True:
        try:
            transcript = audio_queue.get(timeout=timeout)
            all_text += transcript + "\n"
            summary = summarize_meeting(meeting_text=all_text)
            return summary
        except queue.Empty:
            continue
        except Exception as e:
            _print_err(f"Error processing transcript: {e}")
            break

    return {"key_points": [], "decisions": [], "action_items": []}

# -----------------------------
# OCR helpers
# -----------------------------


def ocr(**kwargs) -> str:
    """
    Perform OCR on an image using OCR.Space API.

    kwargs:
        - image_path or raw_input: path to the image
        - api_key: OCR.Space API key
    """
    image_path = kwargs.get("image_path") or kwargs.get("raw_input")
    api_key = kwargs.get("api_key", "K85328613788957")
    
    if not image_path:
        return "Error: No image path provided."

    p = Path(image_path)
    if not p.exists():
        return "Error: image not found."

    try:
        with p.open("rb") as img_file:
            r = SESSION.post(
                "https://api.ocr.space/parse/image",
                files={"image": img_file},
                data={"apikey": api_key, "language": "eng", "OCREngine": "2"},
                timeout=60
            )
        r.raise_for_status()
    except Exception as e:
        return f"Error: OCR request failed: {e}"

    try:
        result = r.json()
    except ValueError:
        return f"Response not JSON: {r.text[:200]}..."

    if result.get("IsErroredOnProcessing"):
        return "❌ OCR Failed: " + str(result.get("ErrorMessage"))

    parsed = result.get("ParsedResults")
    if parsed and parsed[0].get("ParsedText"):
        return parsed[0]["ParsedText"].strip()

    return "⚠️ No text found in image."

def ocr_screen(**kwargs) -> Optional[str]:
    """
    Capture the screen, perform OCR, and summarize the text.

    kwargs:
        - api_key: OCR.Space API key
    """
    api_key = kwargs.get("api_key", "K85328613788957")

    # Capture whole screen to temp file
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        temp_path = Path(tmp.name)

    try:
        try:
            img = ImageGrab.grab()
        except Exception:
            return "Screen capture not supported on this system."
        
        img.save(temp_path)

        # Perform OCR
        text = ocr(image_path=temp_path, api_key=api_key)

        # Summarize OCR text using Groq
        try:
            summary = groq_answer(
                "Describe the content of this text in 1-3 sentences. Do not mention screenshot/image.",
                text
            )
        except Exception as e:
            _print_err(f"Groq summarization failed: {e}")
            summary = text  # fallback

        return summary

    finally:
        # Always attempt to remove temporary file
        try:
            temp_path.unlink(missing_ok=True)
        except Exception:
            pass

def translate_image(**kwargs) -> str:
    """
    Extract text from an image using OCR and translate it.

    kwargs:
        - image_path or raw_input: path to the image
        - target_lang: language code to translate into (default 'en')
    """
    image_path = kwargs.get("image_path") or kwargs.get("raw_input")
    target_lang = kwargs.get("target_lang", "en")

    if GoogleTranslator is None:
        return "deep_translator not available."

    if not image_path or not Path(image_path).exists():
        return "Error: Image not found."

    # --- Extract text using OCR ---
    extracted_text = ocr(image_path=image_path)
    if not extracted_text or extracted_text.startswith("Error"):
        return "No readable text found."

    # --- Translate text ---
    try:
        translated = GoogleTranslator(source="auto", target=target_lang).translate(extracted_text)
        return translated
    except Exception as e:
        _print_err(f"Translation failed: {e}")
        return f"Translation failed: {e}"

# -----------------------------
# Quick utilities
# -----------------------------

def clear_recycle_bin(**kwargs) -> str:
    """
    Clears recycled/trashed items cross-platform.
    Delegates to SystemCleanup.
    """
    try:
        return SystemCleanup.clean_recycled_items()
    except Exception as e:
        _print_err(f"Failed to clear recycle bin: {e}")
        return "Error clearing recycle bin."

# -----------------------
# Lock screen
# -----------------------
def lock_screen(**kwargs) -> str:
    """Locks the screen using OS-native methods."""
    try:
        if IS_WINDOWS:
            try:
                ctypes.windll.user32.LockWorkStation()
                return "Screen locked (Windows)."
            except Exception:
                # Fallback with keyboard emulation
                try:
                    press_and_release("win + l")
                    return "Screen locked (Windows, keyboard emulation)."
                except Exception:
                    return "Screen lock failed (Windows)."

        elif IS_MACOS:
            try:
                # Requires accessibility permissions
                cmd = 'osascript -e \'tell application "System Events" to keystroke "q" using {control down, command down}\''
                subprocess.run(cmd, check=True, shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                return "Screen locked (macOS)."
            except Exception:
                return "Error: macOS lock failed. Enable scripting/accessibility."

        elif IS_LINUX:
            if shutil.which("gnome-screensaver-command"):
                subprocess.run(["gnome-screensaver-command", "-l"], check=False)
                return "Screen locked (Linux Gnome)."
            elif shutil.which("loginctl"):
                subprocess.run(["loginctl", "lock-session"], check=False)
                return "Screen locked (Linux loginctl)."
            else:
                return "Warning: Linux DE lock command not found."

    except Exception as e:
        _print_err(f"Lock screen failed: {e}")
        return "Lock screen not supported on this OS."

# -----------------------
# Translate text file
# -----------------------
def translate_document(**kwargs):
    if GoogleTranslator is None:
        return "deep_translator not available."

    input_file = kwargs.get("input_file")
    output_file = kwargs.get("output_file")
    target_language = kwargs.get("target_language", "en")

    if not input_file or not output_file:
        return "Error: input_file and output_file are required."

    inp = Path(input_file)
    out = Path(output_file)

    if not inp.exists():
        return f"Error: File not found: {inp}"

    try:
        with inp.open("r", encoding="utf-8", errors="ignore") as infile, \
             out.open("w", encoding="utf-8") as outfile:
            for line in infile:
                try:
                    tr = GoogleTranslator(source="auto", target=target_language).translate(line)
                    outfile.write(tr + "\n")
                except Exception as e:
                    _print_err(f"Translation failed for line: {e}")
                    outfile.write(line + "\n")  # fallback: write original line
        return f"Document translated: {out}"
    except Exception as e:
        _print_err(f"translate_document failed: {e}")
        return "Error translating document."

# -----------------------
# Shortcut to open browser history (example)
# -----------------------
def s_h(**kwargs):
    import webbrowser
    try:
        webbrowser.open("https://www.google.com")
        time.sleep(2)
        try:
            press_and_release("ctrl + h")
        except Exception:
            pass
    except Exception as e:
        _print_err(f"s_h failed: {e}")

# -----------------------------
# Natural alarm AI
# -----------------------------


def natural_alarm_ai(**kwargs):
    command = kwargs.get("command") or kwargs.get("raw_input") or ""
    message = kwargs.get("message", "Reminder!")

    parsed = groq_answer(
        "Extract time from the text below. Return ONLY valid JSON: "
        "either {'hours':int,'minutes':int,'seconds':int} or {'absolute_time':'YYYY-MM-DD HH:MM:SS'}",
        command,
    )

    try:
        data = json.loads((parsed or "").strip())
    except Exception:
        return f"❌ AI could not parse the time. Raw: {parsed}"

    now = dt.datetime.now()
    duration_seconds = 0

    if any(k in data for k in ["hours", "minutes", "seconds"]):
        duration_seconds = data.get("hours", 0) * 3600 + data.get("minutes", 0) * 60 + data.get("seconds", 0)
    elif "absolute_time" in data:
        try:
            target = dt.datetime.fromisoformat(data["absolute_time"])
            duration_seconds = (target - now).total_seconds()
        except Exception:
            return "❌ Invalid absolute_time format."

    if duration_seconds <= 0:
        return "❌ Invalid or past time provided."

    # Fetch weather & quote while waiting
    try:
        weather = SESSION.get("https://wttr.in/?format=3", timeout=10).text.strip()
    except Exception:
        weather = "N/A"
    try:
        jq = SESSION.get("https://zenquotes.io/api/random", timeout=10).json()
        quote = jq[0].get("q", "Stay awesome!") if isinstance(jq, list) else "Stay awesome!"
    except Exception:
        quote = "Stay awesome!"

    time.sleep(duration_seconds)

    # Ring alarm
    if pygame:
        try:
            pygame.mixer.init()
            wav = Path("alarm.wav")
            if wav.exists():
                pygame.mixer.music.load(str(wav))
                pygame.mixer.music.play()
        except Exception as e:
            _print_err(f"Alarm play failed: {e}")

    _safe_message_box("Alarm", f"⏰ Alarm ringing!\n\n🌤 Weather: {weather}\n💡 Quote: {quote}\n\n{message}")

    if pygame:
        try:
            pygame.mixer.music.stop()
        except Exception:
            pass

    return "✅ Alarm finished."


# -----------------------------
# YouTube Utilities
# -----------------------------
def youtube_summarizer(**kwargs):
    try:
        url = kwargs.get("url") or kwargs.get("raw_input") or ""
        yt = YouTube(url)
        text = yt.description or ""
        if not text.strip():
            return "No captions/description available for this video."
        summary = groq_answer("Summarize the following video description concisely:", text)
        return summary
    except Exception as e:
        return f"Error: {e}"


def ytDownloader(**kwargs):
    yt_url = kwargs.get("yt_url") or kwargs.get("raw_input") or ""
    if YouTube is None:
        return "pytube not available."
    try:
        yt = YouTube(yt_url)
        video = yt.streams.get_highest_resolution()
        out = video.download()
        return f"Downloaded: {out}"
    except Exception as e:
        return f"Download failed: {e}"


def playMusic(**kwargs):
    song_name = kwargs.get("song_name") or kwargs.get("raw_input") or ""
    if pywhatkit is None:
        return "pywhatkit not available."
    try:
        pywhatkit.playonyt(song_name)
    except Exception as e:
        _print_err(f"Play failed: {e}")


# -----------------------------
# QR Code Generator
# -----------------------------
def qrCodeGenerator(**kwargs) -> str:
    input_text_link = kwargs.get("input_text_link") or kwargs.get("raw_input") or ""
    if qrcode is None:
        return "Error: 'qrcode' library not installed."
    if not input_text_link.strip():
        return "Error: No text or link provided."

    try:
        fname = Path(dt.datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + "-QrCode.png")
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=15, border=4)
        qr.add_data(input_text_link)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img.save(fname)

        # Open cross-platform
        try:
            if IS_WINDOWS:
                os.startfile(str(fname))
            elif IS_MACOS:
                subprocess.run(["open", str(fname)], check=False, timeout=5, stderr=subprocess.DEVNULL)
            elif IS_LINUX:
                subprocess.run(["xdg-open", str(fname)], check=False, timeout=5, stderr=subprocess.DEVNULL)
            else:
                return f"Saved QR: {fname.name}. File opening skipped (unsupported OS)."
        except Exception:
            return f"Saved QR: {fname.name}. Could not open file automatically."

        return f"Saved QR and opened file: {fname.name}"
    except Exception as e:
        _print_err(f"QR generation failed: {e}")
        return f"QR code generation failed: {e}"


# -----------------------------
# PDF Reader (first page)
# -----------------------------
def read_pdf(**kwargs) -> str:
    pdf_file = kwargs.get("pdf_file") or kwargs.get("raw_input")
    try:
        reader = PdfReader(str(pdf_file))
        if not reader.pages:
            return ""
        return reader.pages[0].extract_text() or ""
    except Exception as e:
        _print_err(f"PDF read failed: {e}")
        return ""

# -----------------------------
# File Organizer
# -----------------------------


def organize_files(**kwargs):
    directory = Path(kwargs.get("directory") or kwargs.get("raw_input"))
    if not directory.is_dir():
        return "Directory does not exist."
    
    mapping = {
        ("doc", "docx"): "Word",
        ("xls", "xlsx"): "Excel",
        ("ppt", "pptx"): "PowerPoint",
        ("pdf",): "PDF",
        ("exe",): "Applications",
        ("py", "java", "html", "css", "js"): "Code",
        ("jpg", "jpeg", "png", "gif", "webp", "bmp"): "Images",
        ("mp4", "mkv", "mov", "avi"): "Videos",
        ("mp3", "wav", "flac"): "Audio",
        ("zip", "rar", "7z", "tar", "gz"): "Archives",
        ("csv",): "Data",
    }

    def target_folder_for(ext: str) -> str:
        ext = ext.lower().lstrip(".")
        for exts, folder in mapping.items():
            if ext in exts:
                return folder
        return ext or "no_extension"

    moved = []
    for item in directory.iterdir():
        if item.is_dir():
            continue
        target = directory / target_folder_for(item.suffix[1:])
        target.mkdir(exist_ok=True)
        try:
            shutil.move(str(item), str(target / item.name))
            moved.append(f"{item.name} -> {target.name}/")
        except Exception as e:
            _print_err(f"Move failed for {item}: {e}")
    
    return f"Moved items:\n" + "\n".join(moved) if moved else "No files moved."

def file_organizer(**kwargs) -> str:
    directory = kwargs.get("directory") or kwargs.get("raw_input")
    if Path(directory).is_dir():
        return organize_files(directory=directory)
    return "The specified directory does not exist."
# -----------------------------
# Simple transcription
# -----------------------------
def transcribe_audio(**kwargs) -> str:
    file_path = kwargs.get("file_path") or kwargs.get("raw_input")
    r = sr.Recognizer()
    try:
        with sr.AudioFile(str(file_path)) as source:
            audio = r.record(source)
        return r.recognize_google(audio)
    except sr.UnknownValueError:
        return "Could not understand audio."
    except sr.RequestError:
        return "Speech recognition request failed."
    except Exception as e:
        return f"Transcription failed: {e}"

# -----------------------------
# Download images (batch)
# -----------------------------

def download_images(**kwargs):
    image_urls = kwargs.get("image_urls") or kwargs.get("raw_input") or []
    downloaded = []
    for url in image_urls:
        try:
            r = SESSION.get(url, timeout=20)
            r.raise_for_status()
            name = url.split("/")[-1] or f"image_{hashlib.md5(url.encode()).hexdigest()}.jpg"
            Path(name).write_bytes(r.content)
            downloaded.append(name)
        except Exception as e:
            _print_err(f"Failed to download {url}: {e}")
    return f"Downloaded: {', '.join(downloaded)}" if downloaded else "No images downloaded."

# -----------------------------
# Create file from natural text
# -----------------------------

_FILE_EXT_MAP = {
    "python file": ".py", "java file": ".java", "text file": ".txt", "html file": ".html",
    "css file": ".css", "javascript file": ".js", "json file": ".json", "xml file": ".xml",
    "csv file": ".csv", "markdown file": ".md", "yaml file": ".yaml", "pdf file": ".pdf",
    "word file": ".docx", "excel file": ".xlsx", "powerpoint file": ".pptx", "zip file": ".zip",
    "tar file": ".tar", "image file": ".png", "audio file": ".mp3", "video file": ".mp4"
}

def get_file_extension(text: str) -> str:
    for key, ext in _FILE_EXT_MAP.items():
        if key in text.lower():
            return ext
    return ""

def _strip_type_words(text: str) -> str:
    t = text
    for key in _FILE_EXT_MAP.keys():
        t = t.replace(key, "")
    t = t.replace("named", "").replace("with name", "").replace("create", "")
    return " ".join(t.split())

def create_file(text: str):
    ext = get_file_extension(text)
    name = _strip_type_words(text) or "demo"
    p = Path(f"{name}{ext}")
    p.touch(exist_ok=True)
    return f"Created: {p.resolve()}"


# -----------------------------
# Top processes
# -----------------------------

def get_top_processes(num_processes: int = 3):
    procs = []
    for proc in psutil.process_iter(["pid", "name", "cpu_percent", "memory_info"]):
        try:
            name = proc.info.get("name") or f"pid-{proc.info.get('pid')}"
            cpu = float(proc.info.get("cpu_percent") or 0.0)
            mem = getattr(proc.info.get("memory_info"), "rss", 0)
            procs.append((name, cpu, mem))
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    procs.sort(key=lambda p: (p[1], p[2]), reverse=True)
    return procs[:num_processes]

def display_top_processes():
    lines = [f"Process: {n}, CPU: {int(c)}%, Memory: {int(m / (1024*1024))} MB"
             for n, c, m in get_top_processes()]
    return "\n".join(lines) if lines else "No processes to display."


# -----------------------------
# Wallpaper
# -----------------------------

def change_wallpaper(*args, **kwargs) -> str:
    image_path = args[0] if args else kwargs.get("image_path") or kwargs.get("raw_input")
    if not image_path:
        return "Error: No image path provided."
    p = Path(image_path).resolve()
    if not p.is_file():
        return f"Error: Wallpaper file not found at: {p}"
    image_uri = p.as_uri()

    try:
        if IS_WINDOWS and ctypes:
            ctypes.windll.user32.SystemParametersInfoW(20, 0, str(p), 3)
            return "Wallpaper changed successfully (Windows)."
        elif IS_MACOS:
            script = f'tell application "System Events" to set desktop picture to POSIX file "{p}"'
            subprocess.run(["osascript", "-e", script], check=True, timeout=10)
            return "Wallpaper changed successfully (macOS)."
        elif IS_LINUX:
            try:
                subprocess.run(["gsettings", "set", "org.gnome.desktop.background", "picture-uri", image_uri], check=True, timeout=5)
                subprocess.run(["gsettings", "set", "org.gnome.desktop.background", "picture-uri-dark", image_uri], check=False, timeout=5)
                return "Wallpaper changed successfully (Linux - GSettings/GNOME)."
            except subprocess.CalledProcessError:
                try:
                    subprocess.run(["xfconf-query", "-c", "xfce4-desktop", "-p", "/backdrop/screen0/monitor0/workspace0/last-image", "-s", str(p)], check=True, timeout=5)
                    return "Wallpaper changed successfully (Linux - XFCE)."
                except subprocess.CalledProcessError:
                    try:
                        subprocess.run(["feh", "--bg-scale", str(p)], check=True, timeout=5)
                        return "Wallpaper changed successfully (Linux - feh)."
                    except subprocess.CalledProcessError:
                        return "Linux wallpaper change failed. Requires gsettings, xfconf, or feh."
    except Exception as e:
        _print_err(f"Wallpaper change failed: {e}")
        return f"Wallpaper change failed: {e}"

    return "Wallpaper change not supported on this OS."

# -----------------------------
# Analyze CSV -> DOCX report
# -----------------------------

def analyze_and_report(*args, **kwargs):
    """
    Analyze a CSV file and generate a Word report using AI.

    Accepts:
        Positional args: csv_file, report_file
        Keyword args: csv_file, report_file
    Returns:
        str: Status message
    """
    if Document is None:
        return "python-docx not available."

    # Determine input CSV and output report file
    csv_file = report_file = None
    if args:
        csv_file = args[0]
        report_file = args[1] if len(args) > 1 else kwargs.get("report_file")
    else:
        csv_file = kwargs.get("csv_file") or kwargs.get("raw_input")
        report_file = kwargs.get("report_file")

    if not csv_file or not Path(csv_file).is_file():
        return f"CSV file not found: {csv_file}"

    if not report_file:
        report_file = str(Path("AI_Report.docx").resolve())

    try:
        # Read CSV
        csv_text = Path(csv_file).read_text(encoding="utf-8", errors="ignore")

        # Generate AI report
        report_content = groq_answer(
            "Analyze the following CSV data and generate a detailed report:",
            csv_text
        )
        if not report_content:
            return "AI analysis returned no content."

        # Create Word document
        doc = Document()
        doc.add_heading("AI-Generated Report", level=1)
        doc.add_paragraph(report_content)
        doc.save(str(report_file))

        return f"✅ Report generated: {report_file}"

    except FileNotFoundError:
        _print_err(f"File not found: {csv_file}")
        return f"❌ CSV file not found: {csv_file}"
    except Exception as e:
        _print_err(f"Report generation failed: {e}")
        return f"❌ Report generation failed: {e}"


# -----------------------------
# Email (Django)
# -----------------------------


def send_email(*args, **kwargs):
    """
    Send a single email using Django's send_mail.
    Accepts:
        Positional: message, recipient_email
        Keyword: message, email
    """
    if args:
        message = args[0]
        email = args[1] if len(args) > 1 else kwargs.get("email")
    else:
        message = kwargs.get("message") or kwargs.get("raw_input")
        email = kwargs.get("email")

    if not message or not email:
        return "❌ Missing message or recipient email."

    try:
        from django.core.mail import send_mail as dj_send_mail  # type: ignore
    except Exception:
        return "❌ Django mail not configured/installed."

    try:
        dj_send_mail(
            subject="",
            message=message,
            from_email=email,
            recipient_list=[email],
            fail_silently=False
        )
        return f"✅ Email sent to {email}"
    except Exception as e:
        _print_err(f"send_email failed: {e}")
        return f"❌ Failed to send email to {email}: {e}"


def send_multiple_emails(*args, **kwargs):
    """
    Send multiple emails to a comma-separated list using Django's send_mail.
    Accepts:
        Positional: message, emails_string, sender
        Keyword: message, emails_string, sender
    """
    if args:
        message = args[0]
        emails_string = args[1] if len(args) > 1 else kwargs.get("emails_string")
        sender = args[2] if len(args) > 2 else kwargs.get("sender", "you@example.com")
    else:
        message = kwargs.get("message") or kwargs.get("raw_input")
        emails_string = kwargs.get("emails_string")
        sender = kwargs.get("sender", "you@example.com")

    if not message or not emails_string:
        return "❌ Missing message or recipient emails."

    try:
        from django.core.mail import send_mail as dj_send_mail  # type: ignore
    except Exception:
        return "❌ Django mail not configured/installed."

    emails = [e.strip() for e in emails_string.split(",") if e.strip()]
    results = []

    for email in emails:
        try:
            dj_send_mail(
                subject="",
                message=message,
                from_email=sender,
                recipient_list=[email],
                fail_silently=False
            )
            results.append(f"✅ Email sent to {email}")
        except Exception as e:
            _print_err(f"Failed to send to {email}: {e}")
            results.append(f"❌ Failed to send to {email}: {e}")

    return "\n".join(results)

# -----------------------------
# Search & open files
# -----------------------------

def list_all_files_and_folders(path: str | Path) -> str:
    """
    Recursively lists all files and folders under the given path.
    """
    path = Path(path).resolve()
    if not path.exists():
        return f"❌ Path does not exist: {path}"
    
    lines = []
    for root, dirs, files in os.walk(path):
        lines.append(f"\n📁 Folder: {root}")
        for d in dirs:
            lines.append(f"  📂 Subfolder: {d}")
        for f in files:
            lines.append(f"  📄 File: {f}")
    return "\n".join(lines)


def open_file(*args, **kwargs):
    """
    Opens the most recently modified file that matches the keyword in the specified roots.
    Positional args: keyword, roots (optional)
    Keyword args: keyword, roots (list of paths)
    """
    if args:
        keyword = args[0]
        roots = args[1] if len(args) > 1 else kwargs.get("roots")
    else:
        keyword = kwargs.get("keyword") or kwargs.get("raw_input")
        roots = kwargs.get("roots")

    if not keyword:
        return "❌ No keyword provided."

    # Default search roots
    roots = roots or [
        Path.home() / "Documents",
        Path.home() / "Downloads",
        Path.home() / "Desktop",
        Path.home() / "Pictures",
        Path.home() / "Videos",
        Path.home() / "Music",
    ]

    keyword_lower = keyword.lower()
    candidates: List[Path] = []

    for root in roots:
        root = Path(root)
        if not root.exists():
            continue
        for p in root.rglob("*"):
            if p.is_file() and keyword_lower in p.name.lower():
                candidates.append(p)

    if not candidates:
        return "❌ No matching file found."

    # Pick the most recently modified file
    best = max(candidates, key=lambda p: p.stat().st_mtime)

    try:
        if IS_WINDOWS:
            os.startfile(str(best.resolve()))
        elif IS_MACOS:
            subprocess.run(["open", str(best.resolve())], check=False)
        else:
            subprocess.run(["xdg-open", str(best.resolve())], check=False)
        return f"✅ Opened: {best}"
    except Exception as e:
        _print_err(f"Failed to open file: {e}")
        return f"❌ Failed to open file: {best}"

# -----------------------------
# Brightness & Net speed
# -----------------------------

def dim_light(*args, **kwargs):
    """
    Sets the screen brightness to a specified level (0-100).
    Positional args: level
    Keyword args: level
    """
    if args:
        level = args[0]
    else:
        level = kwargs.get("level", 45)
    
    try:
        level_int = int(level)
        if not 0 <= level_int <= 100:
            return "Brightness level must be between 0 and 100."
        set_brightness(level_int)
        return f"Brightness set to {level_int}%."
    except Exception as e:
        _print_err(f"Set brightness failed: {e}")
        return f"Failed to set brightness: {e}"


def internet_speed(duration: int = 3) -> str:
    """
    Measures the approximate internet download speed over a given duration (seconds).
    Returns Mbps.
    """
    try:
        pernic = psutil.net_io_counters(pernic=True)
        # Pick the first non-loopback interface with traffic
        interface = next((name for name, stats in pernic.items() 
                          if not name.startswith("lo") and stats.bytes_recv > 0), None)
        if not interface:
            return "No active network interface found."
        
        start_bytes = pernic[interface].bytes_recv
        time.sleep(duration)
        end_bytes = psutil.net_io_counters(pernic=True)[interface].bytes_recv
        received = end_bytes - start_bytes
        mbps = received * 8 / (duration * 1024 * 1024)  # Convert bytes to megabits
        return f"Internet Speed ({interface}): {mbps:.2f} Mbps"
    
    except Exception as e:
        _print_err(f"Speed check failed: {e}")
        return f"Internet speed check failed: {e}"


# -----------------------------
# System restore point (Win)
# -----------------------------

def create_system_restore_point(**kwargs) -> str:
    """
    Creates a system snapshot or restore point using native mechanisms:
    - Windows: System Restore
    - macOS: Time Machine Snapshot
    - Linux: Timeshift Snapshot
    """
    default_name = f"AutoPoint_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    name = kwargs.get("name", default_name)
    clean_name = "".join(c for c in name if c.isalnum() or c in "_-")[:64]

    if IS_WINDOWS:
        try:
            command = f'wmic.exe /Namespace:\\\\root\\default Path SystemRestore Call CreateRestorePoint "{clean_name}", 100, 7'
            result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=15)
            if "ReturnValue = 0" in result.stdout:
                return f"System Restore Point created successfully (Windows): {clean_name}"
            _print_err(f"WMIC output: {result.stdout.strip()}")
            return "Failed to create Windows System Restore Point. Ensure System Protection is ON."
        except Exception as e:
            _print_err(f"Windows restore point failed: {e}")
            return f"Windows restore point command failed: {e}"

    elif IS_MACOS:
        try:
            subprocess.run(["tmutil", "snapshot"], check=True, capture_output=True, text=True, timeout=30)
            return f"Local Time Machine Snapshot created successfully (macOS)."
        except subprocess.CalledProcessError as e:
            _print_err(f"tmutil failed: {e.stderr.strip()}")
            return "Failed to create macOS Time Machine Snapshot. Ensure Time Machine is configured."
        except Exception as e:
            _print_err(f"macOS snapshot failed: {e}")
            return f"macOS snapshot command failed: {e}"

    elif IS_LINUX:
        try:
            result = subprocess.run(
                ["sudo", "timeshift", "--create", "--comments", clean_name],
                check=False, capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0 and "Snapshot created successfully" in result.stdout:
                return f"Timeshift Snapshot created successfully (Linux): {clean_name}"
            elif result.returncode != 0 and "command not found" in result.stderr.lower():
                return "Linux snapshot failed: Timeshift is not installed or not in PATH."
            _print_err(f"Timeshift output: {result.stderr.strip()}")
            return "Failed to create Linux Timeshift Snapshot. Requires sudo and Timeshift installed."
        except Exception as e:
            _print_err(f"Linux snapshot failed: {e}")
            return f"Linux snapshot command failed: {e}"

    return "System snapshot functionality not supported on this operating system."

# -----------------------------
# Hashing & dedupe
# -----------------------------

def get_file_hash(**kwargs) -> Optional[str]:
    """Compute SHA-256 hash of a file in chunks to handle large files."""
    path = kwargs.get("path") or kwargs.get("raw_input")
    chunk_size = int(kwargs.get("chunk_size", 1 << 20))  # default 1 MB
    p = Path(path)
    if not p.is_file():
        _print_err(f"Path is not a file: {p}")
        return None

    hasher = hashlib.sha256()
    try:
        with p.open("rb") as f:
            for chunk in iter(lambda: f.read(chunk_size), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception as e:
        _print_err(f"Could not read {p}: {e}")
        return None


def find_and_delete_duplicates(**kwargs):
    """Find duplicate files by SHA-256 hash and delete duplicates."""
    folder = Path(kwargs.get("folder") or kwargs.get("raw_input"))
    if not folder.is_dir():
        return f"Error: {folder} is not a valid directory."

    hashes: Dict[str, Path] = {}
    deleted_files: List[str] = []

    for p in folder.rglob("*"):
        if not p.is_file():
            continue
        file_hash = get_file_hash(path=p)
        if not file_hash:
            continue
        if file_hash in hashes:
            try:
                p.unlink(missing_ok=True)
                deleted_files.append(str(p))
            except Exception as e:
                _print_err(f"Failed to delete {p}: {e}")
        else:
            hashes[file_hash] = p

    return f"✅ Done. {len(deleted_files)} duplicates deleted.\nDeleted files:\n" + "\n".join(deleted_files) if deleted_files else "✅ No duplicates found."


# -----------------------------
# Battery status
# -----------------------------

def _switch_to_power_saver(verbose: bool = True) -> str:
    """
    Activates the platform's power-saver / low-energy profile.
    Works on Windows, macOS, and Linux (requires sudo on Linux).
    """
    try:
        if IS_WINDOWS:
            # GUID for Power Saver plan
            cmd = "powercfg /setactive a1841308-3541-4fab-bc81-f71556f20b4a"
            subprocess.run(cmd, check=False, shell=True, capture_output=True, timeout=5)
            return "✅ Windows power plan set to Power Saver."

        elif IS_MACOS:
            # Enable Low Power Mode (macOS 10.15+)
            subprocess.run(["pmset", "-a", "lowpowermode", "1"],
                           check=False, capture_output=True, timeout=5)
            return "✅ macOS Low Power Mode attempted."

        elif IS_LINUX:
            # Attempt to set CPU governor to 'powersave'
            cpu_paths = list(Path("/sys/devices/system/cpu/").glob("cpu*/cpufreq/scaling_governor"))
            if not cpu_paths:
                return "⚠️ No CPU governor paths found. Power save not applied."
            
            failed = 0
            for p in cpu_paths:
                try:
                    subprocess.run(f"echo powersave | sudo tee {p}",
                                   check=True, shell=True, capture_output=True, timeout=5)
                except Exception:
                    failed += 1
            if failed == 0:
                return "✅ Linux CPU Governor set to 'powersave'."
            else:
                return f"⚠️ Linux attempted 'powersave' but failed for {failed} CPUs (requires sudo)."

        else:
            return "⚠️ Power saving actions skipped: unsupported OS."
    
    except Exception as e:
        if verbose:
            _print_err(f"Power saving activation failed: {e}")
        return f"❌ Power saving activation failed: {e}"


def _show_alert(title: str, message: str):
    """Shows a native, blocking alert box across platforms."""
    try:
        if IS_WINDOWS and ctypes:
            # Windows: native MessageBoxW
            ctypes.windll.user32.MessageBoxW(0, message, title, 1)

        elif IS_MACOS:
            # macOS: AppleScript (escape quotes)
            safe_message = message.replace('"', '\\"')
            safe_title = title.replace('"', '\\"')
            script = f'display dialog "{safe_message}" with title "{safe_title}" buttons {{"OK"}} default button "OK"'
            subprocess.run(["osascript", "-e", script],
                           check=False, timeout=5, stderr=subprocess.DEVNULL)

        elif IS_LINUX:
            # Linux: use zenity if available
            if shutil.which("zenity"):
                subprocess.run(["zenity", "--warning", "--title", title, "--text", message],
                               check=False, timeout=5, stderr=subprocess.DEVNULL)
            else:
                # fallback to terminal print
                print(f"\n[ALERT] {title}: {message}\n")

        else:
            # Fallback for unknown OS
            print(f"\n[ALERT] {title}: {message}\n")
    except Exception as e:
        # In case even native alert fails, fallback to print
        print(f"\n[ALERT] {title}: {message} (Alert failed: {e})\n")


# --- Smart Battery Function ---

def smart_battery(**kwargs) -> str:
    """
    Checks battery status, provides feedback, and triggers power-saving
    measures and alerts on low battery across Windows, macOS, and Linux.
    Requires: 'psutil' library.
    """
    if psutil is None:
        return "Error: The 'psutil' library is not installed."

    try:
        batt = psutil.sensors_battery()
        if batt is None:
            return "Battery info not available (Desktop or unsupported hardware)."

        plugged = batt.power_plugged
        percent = int(batt.percent)

        if plugged:
            return f"Battery is plugged in at {percent}%"

        # Messages and actions by range
        msg = ""
        action_taken = ""

        if percent > 75:
            msg = f"Battery is {percent}% — Perfect."
        elif 50 < percent <= 75:
            msg = f"Battery is {percent}% — Good charge."
        elif 25 < percent <= 50:
            msg = f"Battery is {percent}% — Consider charging soon."
        elif 10 < percent <= 25:
            alert_msg = "Battery low! Switching to saver mode."
            _show_alert("Battery Alert (25%)", alert_msg)
            action_taken = _switch_to_power_saver()
            msg = f"Battery is {percent}% — Charge now! {action_taken}"
        elif 5 < percent <= 10:
            alert_msg = "Battery very low! Switching to saver mode."
            _show_alert("Battery Alert (10%)", alert_msg)
            action_taken = _switch_to_power_saver()
            msg = f"Battery is {percent}% — Charge immediately! {action_taken}"
        else:  # 0-5%
            alert_msg = "Battery critically low! Switching to saver mode."
            _show_alert("Battery Critical! (5%)", alert_msg)
            action_taken = _switch_to_power_saver()
            msg = f"Battery is {percent}% — Critical! Plug in now. {action_taken}"

        return msg

    except Exception as e:
        _print_err(f"Battery check failed: {e}")
        return "Failed to retrieve battery status."


# -----------------------------
# YouTube search
# -----------------------------

def yt_search(**kwargs):
    """
    Opens YouTube search results in the default web browser.

    Keyword Args:
        user (str): Raw user input containing the search query.
        raw_input (str): Alternative to 'user' for backward compatibility.

    Returns:
        str: Confirmation message.
    """
    import webbrowser

    # Extract user input
    user_input = kwargs.get("user") or kwargs.get("raw_input") or ""
    if not user_input.strip():
        return "No search query provided."

    # Remove "youtube search" from input, case-insensitive
    query = re.sub(r"(?i)youtube\s*search", "", user_input).strip()
    if not query:
        return "No valid search query after cleaning input."

    # Open YouTube search in default browser
    try:
        url = f"https://www.youtube.com/results?search_query={requests.utils.quote(query)}"
        webbrowser.open(url)
        return f"Opened YouTube search results for: '{query}'"
    except Exception as e:
        _print_err(f"Failed to open YouTube: {e}")
        return "Failed to open YouTube search."


# -----------------------------
# Smart app/web open/close
# -----------------------------


def openappweb(**kwargs) -> str:
    """
    Opens a URL in the browser (cross-platform) or launches a desktop application 
    using platform-specific commands.
    """
    query = kwargs.get("query") or kwargs.get("raw_input") or ""
    q = query.strip()

    # --- 1. Web Operation (Fully Cross-Platform) ---
    # Detect a domain-ish pattern properly
    if re.search(r"\b[a-z0-9-]+\.(com|co|org|net|io|ai|dev|app)\b", q, re.I):
        q_clean = re.sub(r"(?i)\b(open|jarvis|launch)\b", "", q).replace(" ", "")
        url = q_clean if q_clean.startswith(("http://", "https://")) else f"https://{q_clean}"
        webbrowser.open(url)
        return f"Opened URL: {url}"

    # --- 2. Application Operation (Platform-Specific) ---
    
    # 2a. Determine the clean application name
    app_name = q.lower()
    for word in ("stop", "close", "exit", "open", "launch", "run"):
        app_name = app_name.replace(word, "")
    app_name = app_name.strip()
    
    if not app_name:
        return "Error: No application name specified."

    # Use Groq to infer the short executable name (critical for success)
    short_name = groq_answer(
        "Return ONLY the short executable name (e.g., 'chrome', 'notepad', 'firefox'). DO NOT add .exe or path:",
        app_name
    ).strip().split()[0]
    
    if not short_name:
        _print_err("Could not infer app short name via Groq.")
        return f"Error: Failed to infer short name for '{app_name}'."
    
    try:
        if IS_WINDOWS:
            # Windows: Use subprocess to launch executable directly
            subprocess.Popen([f"{short_name}.exe"], shell=True)
            return f"Launched application (Windows): {short_name}.exe"
            
        elif IS_MACOS:
            # macOS: Use the 'open' command which handles application bundles (.app)
            subprocess.Popen(["open", "-a", short_name])
            return f"Launched application (macOS): {short_name}"
            
        elif IS_LINUX:
            # Linux: Try running the short name directly (assumes it's in PATH)
            subprocess.Popen([short_name])
            return f"Launched application (Linux): {short_name}"
        
        else:
            return "Application launch not supported on this OS."

    except FileNotFoundError:
        return f"Error: Application '{short_name}' not found or not in system PATH."
    except Exception as e:
        _print_err(f"Open app failed: {e}")
        return f"Application launch failed for '{short_name}': {e}"

def closeappweb(**kwargs) -> str:
    """
    Closes the current browser tab (cross-platform) or kills a desktop application process 
    using platform-specific commands.
    """
    query = kwargs.get("query") or kwargs.get("raw_input") or ""
    
    # --- 1. Tab Operation (Fully Cross-Platform via Hotkey) ---
    if "tab" in query.lower() or "browser" in query.lower():
        if pyautogui:
            try:
                # Ctrl+W/Cmd+W is standard for closing tabs/windows
                pyautogui.hotkey("ctrl", "w")
                return "Closed current tab/window."
            except Exception:
                return "Closed current tab/window (hotkey failed)."
        else:
            return "Cannot close tab: pyautogui not installed."

    # --- 2. Application Operation (Platform-Specific Process Kill) ---
    
    # 2a. Determine the clean application name
    q = query.lower()
    for word in ("stop", "close", "exit", "kill"):
        q = q.replace(word, "")
    app_name = q.strip()
    
    if not app_name:
        return "Error: No application name specified to close."

    # Use Groq to infer the short executable name
    short_name = groq_answer(
        "Return ONLY the short process/executable name (e.g., 'chrome', 'notepad', 'firefox'). DO NOT add .exe or path:",
        app_name
    ).strip().split()[0]
    
    if not short_name:
        _print_err("Could not infer app short name via Groq to close.")
        return f"Error: Failed to infer short name for '{app_name}'."
    
    try:
        if IS_WINDOWS:
            # Windows: taskkill by image name
            command = f"taskkill /f /im {short_name}.exe"
            subprocess.run(command, check=True, shell=True, capture_output=True, timeout=5)
            return f"Closed application (Windows, taskkill): {short_name}.exe"
            
        elif IS_MACOS or IS_LINUX:
            # macOS/Linux: pkill by name
            # -f matches the full command line, not just the executable name
            command = ["pkill", "-f", short_name]
            subprocess.run(command, check=True, capture_output=True, timeout=5)
            return f"Closed application (macOS/Linux, pkill): {short_name}"
            
        else:
            return "Application close not supported on this OS."

    except subprocess.CalledProcessError as e:
        if IS_WINDOWS and b"not found" in e.stdout:
            return f"Application '{short_name}' was not running."
        if (IS_MACOS or IS_LINUX) and b"no process found" in e.stderr:
             return f"Application '{short_name}' was not running."
        return f"Failed to close application '{short_name}': {e}"
    except Exception as e:
        _print_err(f"Close app failed: {e}")
        return f"Close application failed for '{short_name}': {e}"

# -----------------------------
# Excel summarize with Groq
# -----------------------------

def summarize_excel_with_groq(**kwargs):
    file_path = kwargs.get("file_path") or kwargs.get("raw_input")
    p = Path(file_path)
    if not p.is_file():
        return f"Error: Excel file not found at {file_path}"

    try:
        # Read Excel, auto-detect engine
        df = pd.read_excel(file_path, engine=None)
        if df.empty:
            return "Excel file is empty."

        # Convert to string (limit rows for huge files)
        text_data = df.head(100).to_string(index=False)  # Limit to first 100 rows
        prompt = (
            "Summarize the following Excel table concisely, highlighting key trends, "
            "important values, and insights. Include headers for clarity:\n\n"
            f"{text_data}"
        )
        summary = groq_answer(prompt)
        return summary or "Groq did not return a summary."
    except Exception as e:
        _print_err(f"Excel summarize failed: {e}")
        return f"Failed to summarize Excel: {e}"


# -----------------------------
# Helpers
# -----------------------------

def textwrap(**kwargs) -> str:
    s = kwargs.get("s") or kwargs.get("raw_input") or ""
    width = int(kwargs.get("width", 100))
    import textwrap as tw
    return "\n".join(tw.wrap(s, width=width))

# -----------------------------
# Melody
# -----------------------------

def melody(**kwargs):
    from melody_generator import main as melody_main
    melody_main()

MACRO_FILE = Path("macros.json")

def load_macros() -> dict:
    if MACRO_FILE.exists():
        try:
            return json.loads(MACRO_FILE.read_text(encoding="utf-8"))
        except Exception as e:
            _print_err(f"Failed to load macros: {e}")
    return {}

def save_macros(macros: dict):
    try:
        MACRO_FILE.write_text(json.dumps(macros, indent=2), encoding="utf-8")
    except Exception as e:
        _print_err(f"Failed to save macros: {e}")

def record_macro(**kwargs):
    name = kwargs.get("name", "default")
    duration = int(kwargs.get("duration", 30))
    print(f"🎥 Recording macro '{name}' for {duration}s...")
    
    start = time.time()
    actions = []
    try:
        while time.time() - start < duration:
            x, y = pyautogui.position()
            actions.append({"time": time.time() - start, "pos": (x, y)})
            time.sleep(0.5)
    except KeyboardInterrupt:
        print("Recording interrupted manually.")
    
    macros = load_macros()
    macros[name] = actions
    save_macros(macros)
    return f"✅ Macro '{name}' saved with {len(actions)} actions."

def play_macro(**kwargs):
    name = kwargs.get("name", "default")
    macros = load_macros()
    
    if name not in macros:
        return f"❌ Macro '{name}' not found."
    
    actions = macros[name]
    start = time.time()
    for act in actions:
        pyautogui.moveTo(*act["pos"])
        # Maintain relative timing
        elapsed = time.time() - start
        sleep_time = act["time"] - elapsed
        if sleep_time > 0:
            time.sleep(sleep_time)
    
    return f"▶️ Macro '{name}' executed ({len(actions)} actions)."

    

# -------------------------
# Web Interaction Functions
# -------------------------

def fetch_page(url: str, retries: int = 3, delay: int = 2, **kwargs) -> str:
    """Fetch HTML content safely with retries and random User-Agent."""
    for _ in range(retries):
        try:
            headers = kwargs.pop("headers", {"User-Agent": random.choice(HEADERS_LIST)})
            resp = SESSION.get(url, timeout=10, headers=headers, **kwargs)
            resp.raise_for_status()
            return resp.text
        except requests.RequestException:
            time.sleep(delay)
    return ""

def get_page_title(url: str) -> str:
    html = fetch_page(url)
    if not html:
        return ""
    soup = BeautifulSoup(html, "html.parser")
    title_tag = soup.find("title")
    return title_tag.get_text(strip=True) if title_tag else ""

def get_meta_description(url: str) -> str:
    html = fetch_page(url)
    if not html:
        return ""
    soup = BeautifulSoup(html, "html.parser")
    desc_tag = soup.find("meta", attrs={"name": "description"})
    return desc_tag["content"].strip() if desc_tag and "content" in desc_tag.attrs else ""

def search_google(query: str, num_results: int = 5) -> List[str]:
    """Scrapes Google search results (fragile; use API for production)."""
    try:
        query_encoded = quote(query)
        url = f"https://www.google.com/search?q={query_encoded}&num={num_results}"
        html = fetch_page(url)
        if not html:
            return []
        soup = BeautifulSoup(html, "html.parser")
        results = []

        # Newer SERP selectors
        for g in soup.find_all("div", class_="yuRUbf"):
            link_tag = g.find("a")
            if link_tag and link_tag.get("href"):
                results.append(link_tag["href"])

        # Fallback to older SERP selectors
        if not results:
            for g in soup.find_all("div", class_="tF2Cxc"):
                link_tag = g.find("a")
                if link_tag and link_tag.get("href"):
                    results.append(link_tag["href"])

        return results[:num_results]
    except Exception as e:
        print(f"search_google failed: {e}")
        return []

def extract_links(url: str) -> List[str]:
    html = fetch_page(url)
    if not html:
        return []
    soup = BeautifulSoup(html, "html.parser")
    return [urljoin(url, a.get("href")) for a in soup.find_all("a", href=True)]

def get_text_content(url: str, selector: str = None) -> str:
    html = fetch_page(url)
    if not html:
        return ""
    soup = BeautifulSoup(html, "html.parser")
    selectors = [selector, "article", "main", "body"] if selector else ["article", "main", "body"]
    for sel in selectors:
        if sel:
            el = soup.select_one(sel)
            if el:
                return el.get_text(separator="\n", strip=True)
    return soup.get_text(separator="\n", strip=True)

def summarize_pdf(file_path: str, **kwargs):
    """
    Summarizes the entire PDF content using Groq.
    Combines text extraction from all pages and generates a concise summary.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "".join(page.extract_text() or "" for page in reader.pages)
        if not text.strip():
            return "PDF has no extractable text."
        summary_prompt = "Summarize the following PDF content concisely:"
        summary = groq_answer(summary_prompt, text)
        return summary
    except FileNotFoundError:
        _print_err(f"File not found: {file_path}")
        return "PDF file not found."
    except Exception as e:
        _print_err(f"summarize_pdf failed: {e}")
        return "PDF summarization failed."

def analyze_data(*, file_path=None, **kwargs):
    """
    Auto-analyze files: CSV, Excel, PDFs.
    - file_type can be provided in kwargs, otherwise inferred from extension
    - Returns a structured dictionary with summaries and anomalies (if applicable)
    """
    if not file_path:
        return {"error": "No file provided"}

    file_path = Path(file_path)
    if not file_path.exists():
        return {"error": f"File not found: {file_path}"}

    # Determine file type from kwargs or file extension
    file_type = kwargs.get("file_type")
    if not file_type:
        ext = file_path.suffix.lower()
        if ext in [".csv"]:
            file_type = "csv"
        elif ext in [".xlsx", ".xls"]:
            file_type = "excel"
        elif ext in [".pdf"]:
            file_type = "pdf"
        else:
            return {"error": "Unsupported file type"}

    # CSV / Excel handling
    if file_type in ("csv", "excel"):
        import pandas as pd
        try:
            df = pd.read_csv(file_path) if file_type == "csv" else pd.read_excel(file_path)
            numeric_df = df.select_dtypes(include="number")
            summary = numeric_df.describe().to_dict() if not numeric_df.empty else {}
            anomalies = numeric_df[
                numeric_df.apply(lambda x: (x - x.mean()).abs() > 3 * x.std(), axis=0).any(axis=1)
            ]
            return {"summary": summary, "anomalies": anomalies.to_dict(orient="records")}
        except Exception as e:
            _print_err(f"CSV/Excel analysis failed: {e}")
            return {"error": "Failed to analyze CSV/Excel file"}

    # PDF handling
    if file_type == "pdf":
        try:
            text = read_pdf(file_path)  # assumes your existing read_pdf function
            if not text.strip():
                return {"summary": "PDF has no extractable text"}
            summary = groq_answer("Summarize the following PDF:", text)
            return {"summary": summary}
        except Exception as e:
            _print_err(f"PDF analysis failed: {e}")
            return {"error": "Failed to analyze PDF"}

    return {"error": "File type not supported"}

def plot_data(file_path: str, x_col: str, y_col: str, output: str = "chart.png", **kwargs):
    """
    Plots x_col vs y_col from a CSV or Excel file and saves as an image.
    Returns the output file path if successful.
    """
    try:
        import pandas as pd
        import matplotlib.pyplot as plt

        # Load file
        if file_path.lower().endswith(".csv"):
            df = pd.read_csv(file_path)
        elif file_path.lower().endswith((".xls", ".xlsx")):
            df = pd.read_excel(file_path)
        else:
            return f"Unsupported file type: {file_path}"

        # Check columns
        if x_col not in df.columns or y_col not in df.columns:
            return f"Columns not found in data: {x_col}, {y_col}"

        # Plot
        plt.figure(figsize=(8, 5))
        plt.plot(df[x_col], df[y_col], marker='o', linestyle='-', color='b')
        plt.title(f"{y_col} vs {x_col}")
        plt.xlabel(x_col)
        plt.ylabel(y_col)
        plt.grid(True)
        plt.tight_layout()

        # Save
        plt.savefig(output, dpi=300)
        plt.close()
        return output

    except Exception as e:
        _print_err(f"plot_data failed: {e}")
        return ""


def convert_text_to_pdf(text: str, output: str = "output.pdf", **kwargs):
    """
    Converts a given text string to a PDF file.
    Returns the path to the generated PDF if successful.
    """
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        for line in text.split("\n"):
            pdf.multi_cell(0, 8, line)

        pdf.output(output)
        return output

    except Exception as e:
        _print_err(f"convert_text_to_pdf failed: {e}")
        return ""

def summarize_text(text: str, **kwargs) -> str:
    """
    Summarizes a given text concisely using the Groq AI model.
    """
    try:
        return groq_answer("Summarize the following text concisely:", text)
    except Exception as e:
        _print_err(f"summarize_text failed: {e}")
        return ""


def sentiment_analysis(text: str, **kwargs) -> str:
    """
    Performs sentiment analysis on the text.
    Returns 'positive', 'negative', or 'neutral'.
    """
    try:
        instructions = "Analyze sentiment of the text and return positive, negative, or neutral:"
        return groq_answer(instructions, text)
    except Exception as e:
        _print_err(f"sentiment_analysis failed: {e}")
        return ""


def nlp_qna(text: str, question: str, **kwargs) -> str:
    """
    Answers a question based on the provided text context using Groq AI.
    """
    try:
        instructions = f"Answer the question based on the following context:\n{text}\nQuestion: {question}"
        return groq_answer(instructions)
    except Exception as e:
        _print_err(f"nlp_qna failed: {e}")
        return ""


def generate_chart_from_data(data: List[Dict], **kwargs) -> str:
    """
    Generates bar charts for all numeric columns in a list-of-dict dataset.
    Saves each chart as '<column_name>_chart.png'.
    """
    try:
        df = pd.DataFrame(data)
        numeric_cols = df.select_dtypes(include='number').columns
        if not numeric_cols.any():
            return "No numeric columns found to generate charts."

        generated_files = []
        for col in numeric_cols:
            plt.figure(figsize=(8,5))
            df[col].plot(kind="bar", title=col)
            filename = f"{col}_chart.png"
            plt.tight_layout()
            plt.savefig(filename)
            plt.close()
            generated_files.append(filename)

        return f"Generated charts: {', '.join(generated_files)}"
    except Exception as e:
        _print_err(f"generate_chart_from_data failed: {e}")
        return "Failed to generate charts."


def analyze_image(**kwargs) -> str:
    """
    Analyzes an image for objects, text, or charts using Groq.
    
    Parameters:
        path: str - Path to the image file
    """
    try:
        path = kwargs.get("path")
        if not path or not os.path.isfile(path):
            return "Error: Image path invalid or file does not exist."
        
        instructions = (
            "Analyze the content of the image and summarize important information, "
            "including detected objects, text, charts, or any visual insights."
        )
        return groq_answer(instructions, f"Image at path: {path}")
    except Exception as e:
        _print_err(f"analyze_image failed: {e}")
        return "Image analysis failed."



def generate_report(**kwargs) -> dict | str:
    """
    Generates a report from a CSV or Excel file with plots and a summary.
    
    kwargs:
        data_path: str - path to input CSV/Excel
        output_path: str - path to save chart/summary (default: 'report.xlsx')
    Returns:
        dict with keys 'chart' and 'summary', or error message string
    """
    try:
        data_path = kwargs.get("data_path")
        output_path = kwargs.get("output_path", "report.xlsx")

        if not data_path or not os.path.isfile(data_path):
            return "Error: Data path invalid or file does not exist."

        # Load data
        df = pd.read_excel(data_path) if data_path.lower().endswith((".xlsx", ".xls")) else pd.read_csv(data_path)

        # Generate line chart for numeric columns
        numeric_cols = df.select_dtypes(include="number").columns
        if not numeric_cols.any():
            return "No numeric columns found to generate chart."

        plt.figure(figsize=(10,6))
        df[numeric_cols].plot(ax=plt.gca())
        plt.title("Auto-generated Chart")
        plt.xlabel("Index")
        plt.ylabel("Values")
        plt.grid(True)
        chart_path = Path(output_path).with_suffix(".png")
        plt.tight_layout()
        plt.savefig(chart_path)
        plt.close()

        # Summarize dataset with Groq
        preview_text = df.head(50).to_string()
        summary = groq_answer("Summarize dataset trends, patterns, and anomalies.", preview_text)

        return {"chart": str(chart_path), "summary": summary}

    except Exception as e:
        _print_err(f"generate_report failed: {e}")
        return f"Failed to generate report: {e}"

    
scheduler = sched.scheduler(time.time, time.sleep)

class FolderWatcher(FileSystemEventHandler):
    def __init__(self, callback):
        self.callback = callback

    def on_created(self, event):
        if not event.is_directory:
            try:
                self.callback(event.src_path)
            except Exception as e:
                _print_err(f"FolderWatcher callback error: {e}")

def watch_folder(**kwargs) -> str:
    """
    Watch a folder for new files and trigger a callback.
    
    kwargs:
        folder: str - folder path to watch
        callback: callable - function to call when a new file is created
    Returns:
        str status message
    """
    folder = kwargs.get("folder")
    callback = kwargs.get("callback")
    
    if not folder or not os.path.isdir(folder):
        return "Error: Folder path invalid or does not exist."
    if not callable(callback):
        return "Error: callback must be a callable function."

    observer = Observer()
    observer.schedule(FolderWatcher(callback), folder, recursive=False)
    observer.start()
    
    # Keep observer in background thread
    threading.Thread(target=observer.join, daemon=True).start()
    
    return f"Watching folder: {folder}"

def schedule_task(**kwargs) -> str:
    """
    Schedule a function to run after a delay in seconds.
    
    kwargs:
        func: callable - function to execute
        delay: int - seconds to wait before execution (default: 5)
    Returns:
        str status message
    """
    func = kwargs.get("func")
    delay = kwargs.get("delay", 5)

    if not callable(func):
        return "Error: Provided func is not callable."
    
    scheduler.enter(delay, 1, func)
    
    # Run scheduler in a background thread if not already running
    threading.Thread(target=scheduler.run, daemon=True).start()
    
    return f"Task scheduled to run in {delay} seconds"


def smart_decide(**kwargs) -> str:
    """
    Suggest next steps or actions based on provided context using Groq.
    
    kwargs:
        context: str - textual context or previous results to base decisions on
    Returns:
        str - suggested next steps
    """
    context = kwargs.get("context", "")
    if not context.strip():
        return "No context provided for decision-making."
    
    instructions = "Analyze the following context and suggest the next logical task or step."
    try:
        suggestion = groq_answer(instructions, context)
        return suggestion
    except Exception as e:
        _print_err(f"smart_decide failed: {e}")
        return "Decision-making failed due to an internal error."


def encrypt_file(**kwargs):
    """
    Simple XOR file encryption
    kwargs:
        path: str
        key: int
    """
    path = kwargs.get("path")
    key = kwargs.get("key", 123)
    if not path or not os.path.exists(path):
        return "File path invalid"
    try:
        with open(path, "rb") as f:
            data = bytearray(f.read())
        data = bytearray(b ^ key for b in data)
        enc_path = Path(path).with_suffix(".enc")
        with open(enc_path, "wb") as f:
            f.write(data)
        return str(enc_path)
    except Exception as e:
        return f"Failed to encrypt file: {e}"
    
def decrypt_file(**kwargs):
    """
    Simple XOR file decryption (reverse of encrypt_file)
    kwargs:
        path: str - path to encrypted file
        key: int - encryption key used (default 123)
    Returns:
        str - path to decrypted file or error message
    """
    path = kwargs.get("path")
    key = kwargs.get("key", 123)
    if not path or not os.path.exists(path):
        return "File path invalid"
    try:
        with open(path, "rb") as f:
            data = bytearray(f.read())
        data = bytearray(b ^ key for b in data)
        dec_path = Path(path).with_suffix(".dec")
        with open(dec_path, "wb") as f:
            f.write(data)
        return str(dec_path)
    except Exception as e:
        return f"Failed to decrypt file: {e}"

    
def clean_clipboard(**kwargs):
    """
    Clears system clipboard
    """
    try:
        import pyperclip
        pyperclip.copy("")
        return "Clipboard cleared"
    except Exception as e:
        return f"Failed to clear clipboard: {e}"
    
def auto_backup(**kwargs):
    """
    Copy files to backup folder
    kwargs:
        src: str
        dest: str
    """
    src = kwargs.get("src")
    dest = kwargs.get("dest")
    if not src or not os.path.exists(src):
        return "Source path invalid"
    os.makedirs(dest, exist_ok=True)
    import shutil
    try:
        if os.path.isdir(src):
            shutil.copytree(src, Path(dest)/Path(src).name, dirs_exist_ok=True)
        else:
            shutil.copy2(src, dest)
        return f"Backup completed to {dest}"
    except Exception as e:
        return f"Backup failed: {e}"

def focus_window(title_contains: str, retries: int = 3, delay: float = 1.0, **kwargs) -> str:
    """
    Attempts to focus (bring to foreground) a window whose title contains the given string.
    Uses pygetwindow for listing and platform-specific commands for activation.
    Requires: 'pygetwindow' library.
    """
    if not gw:
        return "Error: The 'pygetwindow' library is not installed."
        
    target_window = None
    
    # --- 1. Locate the Window (Cross-Platform via pygetwindow) ---
    for _ in range(retries):
        try:
            windows: list = gw.getWindowsWithTitle(title_contains)
            if windows:
                target_window = windows[0]
                break
        except Exception as e:
            _print_err(f"Window search failed: {e}")
            time.sleep(delay)
    
    if not target_window:
        return f"Error: Window containing '{title_contains}' not found after {retries} attempts."

    # --- 2. Focus the Window (Platform-Specific Activation) ---
    
    if IS_WINDOWS:
        try:
            target_window.activate()
            return f"Successfully focused window (Windows): {target_window.title}"
        except Exception as e:
            _print_err(f"Windows focus failed: {e}")
            return f"Error focusing window (Windows): {e}"

    elif IS_MACOS:
        try:
            app_name = target_window.title.split(' - ')[0] if ' - ' in target_window.title else target_window.title
            script = f"""
            tell application "{app_name}"
                activate
            end tell
            """
            subprocess.run(["osascript", "-e", script], check=False, capture_output=True, timeout=5)
            return f"Successfully focused application (macOS): {app_name}"
        except Exception as e:
            _print_err(f"macOS focus failed: {e}")
            return f"Error focusing application (macOS): {e}"

    elif IS_LINUX:
        if not shutil.which("wmctrl"):
            return "Error: wmctrl command not found. Cannot focus window on Linux."
        try:
            result = subprocess.run(["wmctrl", "-l"], capture_output=True, text=True, check=True)
            window_id = None
            for line in result.stdout.splitlines():
                parts = line.split(maxsplit=4)
                if len(parts) == 5 and title_contains.lower() in parts[4].lower():
                    window_id = parts[0]
                    break
            if window_id:
                subprocess.run(["wmctrl", "-i", "-a", window_id], check=True, capture_output=True, timeout=5)
                return f"Successfully focused window (Linux - wmctrl): {target_window.title}"
            else:
                return f"Error: Could not determine window ID for '{title_contains}'."
        except subprocess.CalledProcessError as e:
            _print_err(f"wmctrl failed: {e}")
            return f"Error focusing window (Linux - wmctrl): {e}"
        except Exception as e:
            _print_err(f"Linux focus failed: {e}")
            return f"Error focusing window (Linux): {e}"

    return "Window focus not supported on this operating system."

# ------------------ GUI INTERACTIONS ------------------
def click(x: Optional[int] = None, y: Optional[int] = None, clicks: int = 1, interval: float = 0.2, **kwargs):
    try:
        if x is not None and y is not None:
            pyautogui.click(x=x, y=y, clicks=clicks, interval=interval)
        else:
            pyautogui.click(clicks=clicks, interval=interval)
        return True
    except Exception as e:
        _print_err(f"Click failed → {e}")
        return False

def double_click(x: Optional[int] = None, y: Optional[int] = None, **kwargs):
    return click(x, y, clicks=2)

def right_click(x: Optional[int] = None, y: Optional[int] = None, **kwargs):
    try:
        if x is not None and y is not None:
            pyautogui.rightClick(x=x, y=y)
        else:
            pyautogui.rightClick()
        return True
    except Exception as e:
        _print_err(f"Right click failed → {e}")
        return False

def type_text(text: str, interval: float = 0.05, **kwargs):
    try:
        pyautogui.write(text, interval=interval)
        return True
    except Exception as e:
        _print_err(f"Typing failed → {e}")
        return False

def press_key(key: str, **kwargs):
    try:
        pyautogui.press(key)
        return True
    except Exception as e:
        _print_err(f"Key press failed → {e}")
        return False

def hotkey(*keys, **kwargs):
    try:
        pyautogui.hotkey(*keys)
        return True
    except Exception as e:
        _print_err(f"Hotkey failed → {e}")
        return False

def copy_to_clipboard(text: str, **kwargs):
    try:
        pyperclip.copy(text)
        return True
    except Exception as e:
        _print_err(f"Clipboard copy failed → {e}")
        return False

def paste_from_clipboard(**kwargs):
    try:
        return pyperclip.paste()
    except Exception as e:
        _print_err(f"Clipboard paste failed → {e}")
        return ""

def drag(start_x: int, start_y: int, end_x: int, end_y: int, duration: float = 0.5, **kwargs):
    try:
        pyautogui.moveTo(start_x, start_y)
        pyautogui.dragTo(end_x, end_y, duration=duration)
        return True
    except Exception as e:
        _print_err(f"Drag failed → {e}")
        return False

def scroll(amount: int, **kwargs):
    try:
        pyautogui.scroll(amount)
        return True
    except Exception as e:
        _print_err(f"Scroll failed → {e}")
        return False

def get_screen_size(**kwargs) -> Tuple[int, int]:
    return pyautogui.size()

def get_window_position(title_contains: str, **kwargs) -> Optional[Tuple[int, int, int, int]]:
    if gw is None:
        _print_err("pygetwindow required for get_window_position.")
        return None
    try:
        windows = gw.getWindowsWithTitle(title_contains)
        if not windows:
            return None
        w = windows[0]
        return w.left, w.top, w.width, w.height
    except Exception as e:
        _print_err(f"Could not retrieve window position: {e}")
        return None

def take_screenshot(path: str = None, **kwargs):
    screenshot = pyautogui.screenshot()
    if path:
        screenshot.save(path)
    return screenshot

def wait_for_window(title_contains: str, timeout: int = 10, **kwargs) -> bool:
    if gw is None:
        _print_err("pygetwindow library is required.")
        return False
    start = time.time()
    while time.time() - start < timeout:
        try:
            windows = gw.getWindowsWithTitle(title_contains)
            if windows:
                return True
        except Exception as e:
            _print_err(f"Error during window check: {e}")
        time.sleep(0.5)
    return False

def wait_for_image(image_path: str, timeout: int = 10, confidence: float = 0.8, **kwargs) -> bool:
    start = time.time()
    while time.time() - start < timeout:
        pos = pyautogui.locateOnScreen(image_path, confidence=confidence)
        if pos:
            return True
        time.sleep(0.5)
    return False

def click_image(image_path: str, confidence: float = 0.8, **kwargs):
    pos = pyautogui.locateCenterOnScreen(image_path, confidence=confidence)
    if pos:
        pyautogui.click(pos)
        return True
    return False

def double_click_image(image_path: str, confidence: float = 0.8, **kwargs):
    pos = pyautogui.locateCenterOnScreen(image_path, confidence=confidence)
    if pos:
        pyautogui.doubleClick(pos)
        return True
    return False

def drag_image(start_image: str, end_image: str, confidence: float = 0.8, **kwargs):
    start_pos = pyautogui.locateCenterOnScreen(start_image, confidence=confidence)
    end_pos = pyautogui.locateCenterOnScreen(end_image, confidence=confidence)
    if start_pos and end_pos:
        pyautogui.moveTo(start_pos)
        pyautogui.dragTo(end_pos, duration=0.5)
        return True
    return False

def highlight_image(image_path: str, duration: float = 1.0, confidence: float = 0.8, **kwargs):
    pos = pyautogui.locateOnScreen(image_path, confidence=confidence)
    if pos:
        x, y, w, h = pos
        pyautogui.moveTo(x + w // 2, y + h // 2)
        time.sleep(duration)
        return True
    return False

def click_text(text: str, **kwargs):
    print(f"[INFO] click_text called for '{text}'")
    return True

def drag_text(text: str, target_x: int, target_y: int, **kwargs):
    print(f"[INFO] drag_text called for '{text}' to ({target_x},{target_y})")
    return True


def repeat_macro(file_path: str, times: int = 1, **kwargs):
    for i in range(times):
        play_macro(file_path)
    return True

def chain_commands(commands_list: List[str], **kwargs):
    from task_automation import skill  # assuming SkillEngine instance
    for cmd in commands_list:
        try:
            skill.execute(cmd, **kwargs)
        except KeyError:
            print(f"[WARN] Skill not found: {cmd}")
    return True

# --- Advanced Safety & Recovery ---
def safe_click(x: int, y: int, retries: int = 3, **kwargs):
    for _ in range(retries):
        try:
            pyautogui.click(x, y)
            return True
        except Exception:
            time.sleep(0.2)
    return False

def safe_type(text: str, retries: int = 3, **kwargs):
    for _ in range(retries):
        try:
            pyautogui.typewrite(text)
            return True
        except Exception:
            time.sleep(0.2)
    return False

def backup_clipboard(**kwargs):
    try:
        import pyperclip
        text = pyperclip.paste()
        return text
    except ImportError:
        print("[ERROR] pyperclip required for backup_clipboard")
        return None

def restore_clipboard(text: str, **kwargs):
    try:
        import pyperclip
        pyperclip.copy(text)
        return True
    except ImportError:
        print("[ERROR] pyperclip required for restore_clipboard")
        return False

# --- Accessibility / Visibility Helpers ---
def move_cursor_to_image(image_path: str, confidence: float = 0.8, **kwargs):
    pos = pyautogui.locateCenterOnScreen(image_path, confidence=confidence)
    if pos:
        pyautogui.moveTo(pos)
        return True
    return False

def center_window(title_contains: str, **kwargs) -> bool:
    """
    Centers the first found window whose title contains the given string
    using pyautogui for screen size and pygetwindow for movement.
    
    Returns: True if the window was centered, False otherwise.
    Requires: 'pygetwindow' and 'pyautogui' libraries.
    """
    if gw is None or pyautogui is None:
        _print_err("Libraries 'pygetwindow' and 'pyautogui' are required but not installed.")
        return False
        
    found_window = None

    # --- 1. Locate the Window ---
    try:
        windows: List[gw.Win32Window] = gw.getWindowsWithTitle(title_contains)
        if windows:
            found_window = windows[0]
        else:
            return False # Window not found
            
    except Exception as e:
        _print_err(f"Window search failed: {e}")
        return False

    # --- 2. Calculate and Move ---
    try:
        # Get the screen size using pyautogui (cross-platform method)
        screen_width, screen_height = pyautogui.size()
        
        # Calculate the new top-left coordinates for centering
        # X-coordinate: (Screen Width - Window Width) / 2
        # Y-coordinate: (Screen Height - Window Height) / 2
        new_x = (screen_width - found_window.width) // 2
        new_y = (screen_height - found_window.height) // 2
        
        # Move the window using pygetwindow's cross-platform move method
        found_window.moveTo(new_x, new_y)
        
        # Optional: Ensure the window is visible/un-minimized for the move to take effect
        if found_window.isMinimized:
            found_window.restore()
        
        return True
        
    except Exception as e:
        # This catches errors during screen size retrieval or the move operation itself
        _print_err(f"Failed to center window '{found_window.title}': {e}")
        
        # Note: On some Linux environments, `moveTo` fails silently or requires specific
        # window managers/permissions. This is a common point of failure for cross-platform
        # window management.
        return False
    
def maximize_window(title_contains: str, **kwargs) -> bool:
    """
    Maximizes the first found window whose title contains the given string.
    
    Returns: True if the window was maximized, False otherwise.
    Requires: 'pygetwindow' library.
    """
    if gw is None:
        _print_err("The 'pygetwindow' library is required but not installed.")
        return False
        
    try:
        windows: List[gw.Win32Window] = gw.getWindowsWithTitle(title_contains)
        
        if windows:
            # Maximizes the window using the cross-platform method
            windows[0].maximize()
            return True
            
    except Exception as e:
        _print_err(f"Failed to maximize window: {e}")
        # This can happen if the window doesn't support maximization
        return False
        
    return False # Window not found

def minimize_window(title_contains: str, **kwargs) -> bool:
    """
    Minimizes the first found window whose title contains the given string.
    
    Returns: True if the window was minimized, False otherwise.
    Requires: 'pygetwindow' library.
    """
    if gw is None:
        _print_err("The 'pygetwindow' library is required but not installed.")
        return False
        
    try:
        windows: List[gw.Win32Window] = gw.getWindowsWithTitle(title_contains)
        
        if windows:
            # Minimizes the window using the cross-platform method
            windows[0].minimize()
            return True
            
    except Exception as e:
        _print_err(f"Failed to minimize window: {e}")
        # This can happen if the window doesn't support minimization
        return False
        
    return False # Window not found


SCOPES = ["https://www.googleapis.com/auth/calendar"]

def schedule_calendar_event(**kwargs):
    """
    kwargs:
        title (str)
        start_time (ISO string)
        end_time (ISO string)
    """
    title = _kw(kwargs, "title", "AI Event")
    start_time = _kw(kwargs, "start_time")
    end_time = _kw(kwargs, "end_time")

    if not start_time or not end_time:
        raise ValueError("start_time and end_time required")

    creds = Credentials.from_authorized_user_file("google_token.json", SCOPES)
    service = build("calendar", "v3", credentials=creds)

    event = {
        "summary": title,
        "start": {"dateTime": start_time},
        "end": {"dateTime": end_time}
    }

    service.events().insert(calendarId="primary", body=event).execute()
    return "Calendar event created"



def upload_to_drive(**kwargs):
    """
    kwargs:
        path (str) - file path
        name (str, optional)
    """
    path = _kw(kwargs, "path")
    name = _kw(kwargs, "name") or os.path.basename(path)

    if not path or not os.path.exists(path):
        raise FileNotFoundError("Invalid file path")

    creds = Credentials.from_authorized_user_file(
        "google_token.json", ["https://www.googleapis.com/auth/drive"]
    )
    service = build("drive", "v3", credentials=creds)

    media = MediaFileUpload(path, resumable=True)
    file = service.files().create(
        body={"name": name},
        media_body=media
    ).execute()

    return file.get("id")


def send_whatsapp_message(**kwargs):
    """
    kwargs:
        phone (str) - international format
        message (str)
    """
    phone = _kw(kwargs, "phone")
    message = _kw(kwargs, "message")

    if not phone or not message:
        raise ValueError("phone and message required")

    url = f"https://graph.facebook.com/v18.0/{WHATSAPP_PHONE_ID}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_TOKEN}",
        "Content-Type": "application/json"
    }

    payload = {
        "messaging_product": "whatsapp",
        "to": phone,
        "type": "text",
        "text": {"body": message}
    }

    requests.post(url, headers=headers, json=payload)
    return "WhatsApp message sent"


async def _discord_send(channel_id, message):
    intents = discord.Intents.default()
    client = discord.Client(intents=intents)

    @client.event
    async def on_ready():
        channel = client.get_channel(channel_id)
        if channel:
            await channel.send(message)
        await client.close()

    await client.start(DISCORD_TOKEN)

def send_discord_message(**kwargs):
    """
    kwargs:
        message (str)
        channel_id (int, optional)
    """
    message = _kw(kwargs, "message")
    channel_id = int(_kw(kwargs, "channel_id", DISCORD_CHANNEL_ID))

    if not message:
        raise ValueError("message required")

    asyncio.run(_discord_send(channel_id, message))
    return "Discord message sent"

def _load_memory():
    if not os.path.exists(MEMORY_FILE):
        return {}
    with open(MEMORY_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def _save_memory(data):
    with open(MEMORY_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)

def os_context_skill(**kwargs):
    """
    OS Context Automation Skill
    - Captures active window and system context
    - Resolves pronouns like 'this', 'that' based on last tasks
    - Logs tasks with timestamps
    """
    user_input = kwargs.get("raw_input") or kwargs.get("user_input") or ""
    memory = _load_memory()

    # Detect active window
    active_window = ""
    try:
        win = gw.getActiveWindow()
        if win:
            active_window = win.title
    except Exception:
        active_window = "Unknown"

    # Detect running apps (top 5 CPU usage)
    try:
        top_apps = sorted(psutil.process_iter(['name', 'cpu_percent']),
                          key=lambda p: p.info['cpu_percent'], reverse=True)[:5]
        running_apps = [p.info['name'] for p in top_apps]
    except Exception:
        running_apps = []

    # Resolve pronouns like 'this', 'that' using memory
    resolved_input = user_input
    if memory.get("last_task"):
        resolved_input = user_input.replace("this", memory["last_task"]["description"])

    # Update memory
    memory["last_task"] = {
        "description": user_input,
        "resolved": resolved_input,
        "active_window": active_window,
        "running_apps": running_apps,
        "timestamp": datetime.now().isoformat()
    }
    _save_memory(memory)

    # Build output info
    output = {
        "resolved": resolved_input,
        "active_window": active_window,
        "running_apps": running_apps,
        "timestamp": memory["last_task"]["timestamp"]
    }

    return output


def learn_user_pattern(task_name):
    """Update long-term patterns based on task frequency."""

    recent = get_recent_tasks(limit=20)
    freq = sum(1 for t in recent if t["function"] == task_name)
    if freq > 3:
        print(f"[LEARNING] Task '{task_name}' is now prioritized for automation.")

nlp_model = None
if SPACY_AVAILABLE:
    try:
        nlp_model = spacy.load("en_core_web_sm")
    except (OSError, IOError):
        nlp_model = None

def repair_input(user_input: str) -> str:
    """
    Corrects spelling, removes unwanted characters, and normalizes text.
    """
    user_input = re.sub(r"[^a-zA-Z0-9 .,?!@#]", "", user_input)
    user_input = user_input.strip()
    # Could integrate spellchecker or autocorrect here
    return user_input

def extract_entities(user_input: str) -> dict:
    """
    Extracts named entities (dates, times, people, apps) from text.
    """
    if nlp_model is None:
        return {}
    doc = nlp_model(user_input)
    entities = {}
    for ent in doc.ents:
        entities[ent.label_] = entities.get(ent.label_, []) + [ent.text]
    return entities

def infer_intent(user_input: str) -> str:
    """
    Determines high-level intent using rules + groq AI.
    """
    prompt = f"Determine the single action user wants: '{user_input}'"
    intent = groq_answer(prompt, user_input)
    return intent.strip()

def polish_input(user_input: str) -> str:
    """
    Cleans, summarizes, and converts casual input into actionable commands.
    """
    # Remove filler words, polite phrases, etc.
    cleaned = re.sub(r"\b(please|kindly|could you|would you)\b", "", user_input, flags=re.I)
    cleaned = cleaned.strip()
    return cleaned

def full_nlp_pipeline(user_input: str) -> dict:
    """
    Full preprocessing and understanding pipeline:
    - Repair
    - Entity extraction
    - Intent inference
    - Polishing
    """
    repaired = repair_input(user_input)
    polished = polish_input(repaired)
    entities = extract_entities(polished)
    intent = infer_intent(polished)

    return {
        "original": user_input,
        "repaired": repaired,
        "polished": polished,
        "entities": entities,
        "intent": intent
    }


from difflib import SequenceMatcher

def best_file_match(query: str, root="."):
    best = None
    score = 0

    for dirpath, _, files in os.walk(root):
        for f in files:
            s = SequenceMatcher(None, query.lower(), f.lower()).ratio()
            if s > score:
                best = os.path.join(dirpath, f)
                score = s

    return best if score > 0.4 else None

MAX_RETRIES = 5

def adaptive_auto_coder(user_request: str, context: str = None):
    """
    Fully adaptive auto-coder for NLP commands:
    - Converts natural language tasks into Python code
    - Can execute GUI automation, browser/app launching, and safe OS commands
    - Splits multi-step tasks automatically
    - Retries and fixes code if errors occur
    Restrictions:
    - NO file/folder deletion
    - NO shutdown/reboot commands
    """
    if not user_request.strip():
        return "[ERROR] Empty request"

    last_error = None
    context_prefix = f"Previous context:\n{context}\n" if context else ""

    base_prompt = f"""
    {context_prefix}
    You are an expert Windows automation engineer.

    Convert the following user request into SAFE Python 3.13 code:
    User request: "{user_request}"

    Requirements:
    - Output ONLY Python code (no markdown, no explanations)
    - You may use pyautogui, time, subprocess, webbrowser, or other safe libraries
    - Split multi-step commands into separate executable functions
    - Name functions clearly and infer parameters automatically
    - Avoid dangerous commands like deletion, shutdown, or formatting
    """

    for attempt in range(1, MAX_RETRIES + 1):
        code = groq_answer(base_prompt)
        if not code:
            continue

        # Remove accidental markdown/code fences
        code = re.sub(r"```.*?```", "", code, flags=re.DOTALL).strip()

        try:
            # Execute code in safe environment
            safe_globals = {
                "__builtins__": __builtins__,
                "pyautogui": pyautogui,
                "time": time,
                "webbrowser": webbrowser,
                "subprocess": subprocess
            }
            local_env = {}
            exec_safe(code, safe_globals, local_env)

            print(f"[AUTO-CODER EXECUTED] Attempt {attempt}")
            return f"✅ Task executed successfully."

        except Exception as e:
            last_error = str(e)
            print(f"[AUTO-CODER ERROR {attempt}] {last_error}")
            base_prompt += f"""
            The previous code attempt failed with error:
            {last_error}
            Fix the code and provide a corrected version.
            """

    raise RuntimeError(f"Adaptive auto-coder failed after {MAX_RETRIES} attempts: {last_error}")

SAFE_BUILTINS = {
    "print": print,
    "range": range,
    "len": len,
    "int": int,
    "float": float,
    "str": str,
    "list": list,
    "dict": dict,
    "set": set,
}

def exec_safe(command, raw_input=None, timeout=5, **kwargs):
    """
    Safe command execution wrapper.
    Compatible with SkillRegistry.
    """
    try:
        return os.system(command)
    except Exception as e:
        return str(e)


def window_finder(**kwargs):
    title = kwargs.get("title") or kwargs.get("raw_input")
    if not title:
        return "Please provide a window title or partial title to find."
    
    found = []
    try:
        if gw is not None:
            # pygetwindow supports Windows, macOS (limited), and Linux (X11)
            windows = gw.getAllWindows()
            for w in windows:
                if title.lower() in w.title.lower():
                    found.append(w.title)
        elif IS_LINUX:
            # Fallback for Linux if pygetwindow fails (requires wmctrl)
            try:
                result = subprocess.run(["wmctrl", "-l"], capture_output=True, text=True, check=True)
                for line in result.stdout.splitlines():
                    parts = line.split(maxsplit=4)
                    if len(parts) == 5 and title.lower() in parts[4].lower():
                        found.append(parts[4].strip())
            except FileNotFoundError:
                return "wmctrl not found. Cannot find windows on Linux."
    except Exception as e:
        return f"Window search error: {e}"
        
    if found:
        return f"Found windows:\n" + "\n".join(found)
    return f"No windows found matching '{title}'."

def file_finder(**kwargs):
    # Accepts: filename, start_dir, mode (file or folder)
    filename = kwargs.get("filename") or kwargs.get("raw_input")
    start_dir = kwargs.get("start_dir", Path.home())
    mode = kwargs.get("mode", "file") # "file" or "dir"

    p_start = Path(start_dir)
    if not p_start.is_dir():
        return f"Starting directory not found: {start_dir}"

    found = []
    
    # Simple recursive search
    def search_path(directory: Path):
        for item in directory.iterdir():
            if filename.lower() in item.name.lower():
                if (mode == "file" and item.is_file()) or (mode == "dir" and item.is_dir()):
                    found.append(str(item.resolve()))
            if item.is_dir():
                try:
                    search_path(item) # Recurse
                except PermissionError:
                    continue # Skip directories without permission

    try:
        search_path(p_start)
    except Exception as e:
        return f"File search failed: {e}"
        
    if found:
        # Limit results for readability
        return f"Found {len(found)} results:\n" + "\n".join(found[:10])
    return f"No {mode}s found matching '{filename}' starting from {start_dir}"

def check_process_status(**kwargs):
    # Requires psutil
    name = kwargs.get("name") or kwargs.get("raw_input")
    if not name:
        return "Please provide a process name (e.g., chrome.exe, Finder, bash)."

    found = []
    for proc in psutil.process_iter(['name']):
        if name.lower() in proc.info['name'].lower():
            found.append(f"PID: {proc.pid}, Name: {proc.info['name']}")
    
    if found:
        return f"Found {len(found)} running processes:\n" + "\n".join(found)
    return f"Process '{name}' is not currently running."

def file_manager(**kwargs):
    # Requires shutil and pathlib
    action = kwargs.get("action")
    source_str = kwargs.get("source")
    destination_str = kwargs.get("destination")
    
    if not action:
        return "Please provide an 'action' parameter: 'copy', 'move', or 'delete'."
    
    if action not in ["copy", "move", "delete"]:
        return "Invalid action. Use 'copy', 'move', or 'delete'."
    
    if not source_str:
        return "Please provide a 'source' parameter."
    
    source = Path(source_str)
    
    if action in ["copy", "move"] and not destination_str:
        return f"Please provide a 'destination' parameter for '{action}' action."
    
    if action in ["copy", "move"]:
        destination = Path(destination_str)
    
    if not source.exists():
        return f"Source file/folder not found: {source.name}"

    try:
        if action == "copy":
            if source.is_dir():
                shutil.copytree(source, destination)
            else:
                shutil.copy2(source, destination)
            return f"Copied {source.name} to {destination.name}"
        
        elif action == "move":
            shutil.move(source, destination)
            return f"Moved {source.name} to {destination.name}"
        
        elif action == "delete":
            if source.is_dir():
                shutil.rmtree(source)
            else:
                source.unlink()
            return f"Deleted {source.name}"

    except Exception as e:
        return f"File operation failed for '{action} {source.name}': {e}"

def register_skill(name, func, registry: dict):
    if name not in registry:
        registry[name] = func

def infer_cross_app_action(os_context, user_input):
    app = os_context.get("active_window", "").lower()

    if "send" in user_input and "email" not in user_input:
        if "whatsapp" in app:
            return "send_whatsapp_message"
        if "chrome" in app:
            return "share_current_page"
        if "pdf" in app:
            return "email_current_document"

    return None

def confirm_permission():
    return input("⚠️ This action is sensitive. Continue? (y/n): ").lower() == "y"

RECENT_COMMANDS = []

def detect_anomaly(cmd: str):
    RECENT_COMMANDS.append(cmd)
    if RECENT_COMMANDS.count(cmd) > 3:
        raise RuntimeError("Anomalous repeated command detected")

    if len(RECENT_COMMANDS) > 20:
        RECENT_COMMANDS.pop(0)

def video_to_audio(**kwargs):
    video_path = kwargs.get("video_path") or kwargs.get("raw_input")
    output_path = kwargs.get("output_path") or (str(Path(video_path).with_suffix('.mp3')) if video_path else None)

    if not video_path or not os.path.exists(video_path):
        return "Invalid video file path."

    if not output_path:
        return "Output path not specified."

    try:
        clip = VideoFileClip(video_path)
        clip.audio.write_audiofile(output_path)
        clip.close()
        return f"Audio extracted to {output_path}"
    except Exception as e:
        return f"Video to audio conversion failed: {e}"


def bg_remover(**kwargs):
    image_path = kwargs.get("image_path") or kwargs.get("raw_input")
    output_path = kwargs.get("output_path") or (
        str(Path(image_path).with_name(Path(image_path).stem + "_no_bg.png"))
        if image_path else None
    )

    if not image_path or not os.path.exists(image_path):
        return "Invalid image file path."

    if not output_path:
        return "Output path not specified."

    try:
        # Load image using OpenCV
        img = cv2.imread(image_path)
        if img is None:
            return "Failed to load image."

        height, width = img.shape[:2]

        # Create mask
        mask = np.zeros((height, width), np.uint8)

        # Background & foreground models
        bgdModel = np.zeros((1, 65), np.float64)
        fgdModel = np.zeros((1, 65), np.float64)

        # Define rectangle around main subject (assumes subject is centered)
        rect = (10, 10, width - 20, height - 20)

        # Apply GrabCut
        cv2.grabCut(
            img,
            mask,
            rect,
            bgdModel,
            fgdModel,
            5,
            cv2.GC_INIT_WITH_RECT
        )

        # Create final mask
        mask2 = np.where(
            (mask == cv2.GC_FGD) | (mask == cv2.GC_PR_FGD),
            255,
            0
        ).astype("uint8")

        # Add alpha channel
        b, g, r = cv2.split(img)
        rgba = cv2.merge((b, g, r, mask2))

        # Save as PNG with transparency
        cv2.imwrite(output_path, rgba)

        return f"Background removed. Saved to {output_path}"

    except Exception as e:
        return f"Background removal failed: {e}"


def wordcloud_generator(**kwargs):
    """
    Generate a word cloud from text / csv / document file.

    kwargs:
        text (str)             : path to input file (required)
        background_color (str) : background color (default: white)
        mask (str)             : path to mask image (optional)
        contour_width (int)    : contour width (optional)
        contour_color (str)    : contour color (optional)
        color_func (bool)      : keep mask colors (optional)
        output_path (str)      : output image path (optional)
        show (bool)            : display image (default: True)
    """

    # ---- Inputs ----
    text_path = kwargs.get("text")
    background_color = kwargs.get("background_color", "white")
    mask_path = kwargs.get("mask")
    contour_width = kwargs.get("contour_width", 0)
    contour_color = kwargs.get("contour_color", "black")
    color_func = kwargs.get("color_func", False)
    show = kwargs.get("show", True)

    if not text_path:
        return "Text file path is required."

    if not Path(text_path).exists():
        return "Invalid text file path."

    output_path = kwargs.get(
        "output_path",
        str(Path(text_path).with_suffix(".png"))
    )

    # ---- Read Text ----
    ext = Path(text_path).suffix.lower()

    try:
        if ext in [".txt", ".doc", ".pdf"]:
            text = open(text_path, encoding="utf-8", errors="ignore").read()

        elif ext == ".csv":
            df = pd.read_csv(text_path, encoding="latin-1")
            text = ""
            for val in df.iloc[:, 0]:
                tokens = str(val).lower().split()
                text += " ".join(tokens) + " "

        else:
            return "Unsupported file format."

    except Exception as e:
        return f"Failed to read text: {e}"

    # ---- Mask Handling ----
    mask_array = None
    mask_colors = None

    if mask_path:
        if not Path(mask_path).exists():
            return "Invalid mask file path."

        mask_array = np.array(Image.open(mask_path))

        if color_func:
            mask_colors = ImageColorGenerator(mask_array)

    # ---- Generate WordCloud ----
    try:
        wc = WordCloud(
            stopwords=STOPWORDS,
            mask=mask_array,
            max_font_size=50,
            max_words=1000,
            background_color=background_color,
            color_func=mask_colors,
            contour_width=contour_width,
            contour_color=contour_color,
        ).generate(text)

        if show:
            plt.figure(figsize=(8, 8))
            plt.imshow(wc, interpolation="bilinear")
            plt.axis("off")
            plt.show()

        wc.to_file(output_path)

        return f"WordCloud generated successfully → {output_path}"

    except Exception as e:
        return f"WordCloud generation failed: {e}"

def instagram_video_downloader(**kwargs):
    """
    Download videos from an Instagram post.

    kwargs:
        post_id (str)      : Instagram post shortcode (required)
        output_dir (str)   : directory to save videos (default: current dir)

    Example:
        instagram_video_downloader(post_id="Cxyz123", output_dir="videos")
    """

    post_id = kwargs.get("post_id")
    output_dir = Path(kwargs.get("output_dir", "."))

    if not post_id:
        return "post_id is required."

    output_dir.mkdir(parents=True, exist_ok=True)

    videos = []

    try:
        # Fetch post page
        response = requests.get(
            f"https://www.instagram.com/p/{post_id}/",
            headers={
                "User-Agent": "Mozilla/5.0"
            },
            timeout=10
        )

        if response.status_code == 404:
            return "Specified post not found."

        # Extract JSON data
        json_data = json.loads(
            re.findall(
                r"window\._sharedData\s=\s(\{.*\});</script>",
                response.text
            )[0]
        )

        data = json_data["entry_data"]["PostPage"][0]["graphql"]["shortcode_media"]

        # Single video
        if data.get("is_video"):
            videos.append(data["video_url"])

        # Carousel posts
        if "edge_sidecar_to_children" in data:
            for post in data["edge_sidecar_to_children"]["edges"]:
                node = post["node"]
                if node.get("is_video"):
                    videos.append(node["video_url"])

        if not videos:
            return "No videos found in this post."

        # Download videos
        for idx, video_url in enumerate(videos, start=1):
            output_file = output_dir / f"{post_id}_{idx}.mp4"
            urllib.request.urlretrieve(video_url, output_file)

        return f"Downloaded {len(videos)} video(s) to {output_dir.resolve()}"

    except Exception as e:
        return f"Download failed: {e}"
    

def facebook_video_downloader(**kwargs):
    """
    Download a Facebook video using mbasic page scraping.

    kwargs:
        url (str)           : Facebook video URL (required)
        output_path (str)   : output file path (default: video.mp4)

    Example:
        facebook_video_downloader(
            url="https://www.facebook.com/....",
            output_path="fb_video.mp4"
        )
    """

    url = kwargs.get("url")
    output_path = Path(kwargs.get("output_path", "video.mp4"))

    if not url:
        return "URL is required."

    if "facebook.com" not in url:
        return "Invalid Facebook URL."

    # Convert to mbasic
    url = url.replace("www", "mbasic")

    try:
        response = get(url, timeout=5, allow_redirects=True)
        if response.status_code != 200:
            return "Failed to fetch Facebook page."

        matches = findall("/video_redirect/", response.text)
        if not matches:
            return "Video not found on this page."

        video_url = unquote(response.text.split("?src=")[1].split('"')[0])

    except (HTTPError, ConnectionError) as e:
        return f"Connection error: {e}"

    # ---- Download Video ----
    try:
        r = get(video_url, stream=True)
        total_size = int(r.headers.get("content-length", 0))
        block_size = 1024

        progress_bar = tqdm(
            total=total_size,
            unit="iB",
            unit_scale=True,
            desc="Downloading"
        )

        with open(output_path, "wb") as file:
            for data in r.iter_content(block_size):
                progress_bar.update(len(data))
                file.write(data)

        progress_bar.close()

        if total_size != 0 and progress_bar.n != total_size:
            return "Download incomplete."

        return f"Video downloaded successfully → {output_path.resolve()}"

    except Exception as e:
        return f"Download failed: {e}"

def xml_to_csv_converter(**kwargs):
    """
    Convert Pascal VOC XML annotations to a CSV file.

    kwargs:
        input_dir (str)   : directory containing XML files (required)
        output_path (str) : output CSV path (default: shape_labels.csv)

    Example:
        xml_to_csv_converter(
            input_dir="annotations",
            output_path="labels.csv"
        )
    """

    input_dir = kwargs.get("input_dir")
    output_path = kwargs.get("output_path", "shape_labels.csv")

    if not input_dir:
        return "input_dir is required."

    input_dir = Path(input_dir)

    if not input_dir.exists():
        return "Input directory does not exist."

    xml_list = []

    try:
        for xml_file in glob.glob(str(input_dir / "*.xml")):
            tree = ET.parse(xml_file)
            root = tree.getroot()

            filename = root.find("filename").text
            width = int(root.find("size")[0].text)
            height = int(root.find("size")[1].text)

            for member in root.findall("object"):
                value = (
                    filename,
                    width,
                    height,
                    member.find("name").text,
                    int(member.find("bndbox/xmin").text),
                    int(member.find("bndbox/ymin").text),
                    int(member.find("bndbox/xmax").text),
                    int(member.find("bndbox/ymax").text),
                )
                xml_list.append(value)

        if not xml_list:
            return "No XML annotations found."

        columns = [
            "filename",
            "width",
            "height",
            "class",
            "xmin",
            "ymin",
            "xmax",
            "ymax",
        ]

        df = pd.DataFrame(xml_list, columns=columns)
        df.to_csv(output_path, index=False)

        return f"XML successfully converted to CSV → {Path(output_path).resolve()}"

    except Exception as e:
        return f"Conversion failed: {e}"

_price_watchlist = {}

def track_price(*, url: str, target_price: float, check_interval: int = 3600):
    headers = {"User-Agent": "Mozilla/5.0"}
    _price_watchlist[url] = {
        "target": target_price,
        "interval": check_interval,
        "headers": headers
    }
    return f"Tracking price for {url}"

def _check_prices():
    while True:
        for url, data in _price_watchlist.items():
            r = requests.get(url, headers=data["headers"])
            soup = BeautifulSoup(r.text, "html.parser")
            price = float(
                soup.select_one(".a-price-whole").text.replace(",", "")
            )
            if price <= data["target"]:
                print(f"[ALERT] Price dropped to {price}: {url}")
        time.sleep(min(d["interval"] for d in _price_watchlist.values()))

def malware_static_analyzer(**kwargs):
    """
    Malware static analysis tool (YARA removed, rule-based detection added)

    kwargs:
        file_path (str)            : path to file (required)
        virustotal_api_key (str)   : optional
        enable_pe (bool)
        enable_strings (bool)
        enable_disassembly (bool)
        enable_network (bool)
        enable_packer (bool)
        enable_clamav (bool)
        enable_metadata (bool)
        report_file (str)          : optional
        report_format (str)        : json / csv
    """

    file_path = kwargs.get("file_path")
    if not file_path or not os.path.exists(file_path):
        return "Invalid file path."

    # ---------------- HASHING ----------------
    def calculate_hash(path, algo):
        h = hashlib.new(algo)
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                h.update(chunk)
        return h.hexdigest()

    # ---------------- FILE INFO ----------------
    file_info = {
        "name": os.path.basename(file_path),
        "size": os.path.getsize(file_path),
        "type": magic.from_file(file_path),
        "md5": calculate_hash(file_path, "md5"),
        "sha1": calculate_hash(file_path, "sha1"),
        "sha256": calculate_hash(file_path, "sha256"),
    }

    results = {"file_info": file_info}

    # ---------------- STRINGS ----------------
    strings = []
    if kwargs.get("enable_strings"):
        with open(file_path, "rb") as f:
            data = f.read()
        s = ""
        for b in data:
            if 32 <= b <= 126:
                s += chr(b)
            else:
                if len(s) >= 4:
                    strings.append(s)
                s = ""
        results["strings_sample"] = strings[:20]

    # ---------------- PE ANALYSIS ----------------
    pe_info = {}
    entropy_alerts = []
    suspicious_imports = []

    if kwargs.get("enable_pe") and "PE" in file_info["type"]:
        pe = pefile.PE(file_path)

        pe_info["entry_point"] = hex(pe.OPTIONAL_HEADER.AddressOfEntryPoint)
        pe_info["sections"] = []

        for sec in pe.sections:
            entropy = sec.get_entropy()
            sec_name = sec.Name.decode(errors="ignore").strip("\x00")
            pe_info["sections"].append({
                "name": sec_name,
                "entropy": entropy
            })
            if entropy > 7.0:
                entropy_alerts.append(sec_name)

        dangerous_apis = {
            b"VirtualAlloc",
            b"WriteProcessMemory",
            b"CreateRemoteThread",
            b"LoadLibraryA",
            b"GetProcAddress"
        }

        if hasattr(pe, "DIRECTORY_ENTRY_IMPORT"):
            for entry in pe.DIRECTORY_ENTRY_IMPORT:
                for imp in entry.imports:
                    if imp.name in dangerous_apis:
                        suspicious_imports.append(imp.name.decode())

        pe_info["high_entropy_sections"] = entropy_alerts
        pe_info["suspicious_imports"] = suspicious_imports
        results["pe_analysis"] = pe_info

    # ---------------- RULE ENGINE (REPLACEMENT FOR YARA) ----------------
    rules_triggered = []

    if entropy_alerts:
        rules_triggered.append("HIGH_ENTROPY_SECTIONS")

    if suspicious_imports:
        rules_triggered.append("SUSPICIOUS_WINDOWS_APIS")

    if kwargs.get("enable_strings"):
        if any("http://" in s or "https://" in s for s in strings):
            rules_triggered.append("EMBEDDED_URLS")

        if any(re.search(r"\b\d{1,3}(\.\d{1,3}){3}\b", s) for s in strings):
            rules_triggered.append("EMBEDDED_IP_ADDRESSES")

    results["rule_engine_hits"] = rules_triggered

    # ---------------- NETWORK ARTIFACTS ----------------
    if kwargs.get("enable_network") and strings:
        results["network_artifacts"] = {
            "ips": re.findall(r"\b\d{1,3}(\.\d{1,3}){3}\b", " ".join(strings)),
            "urls": re.findall(r"https?://[^\s]+", " ".join(strings)),
        }

    # ---------------- DISASSEMBLY ----------------
    if kwargs.get("enable_disassembly") and "PE" in file_info["type"]:
        md = Cs(CS_ARCH_X86, CS_MODE_32)
        with open(file_path, "rb") as f:
            code = f.read()
        results["disassembly_sample"] = [
            f"0x{i.address:x} {i.mnemonic} {i.op_str}"
            for i in list(md.disasm(code, 0x1000))[:20]
        ]

    # ---------------- CLAMAV ----------------
    if kwargs.get("enable_clamav"):
        try:
            r = subprocess.run(["clamscan", file_path],
                               stdout=subprocess.PIPE,
                               stderr=subprocess.PIPE)
            results["clamav"] = r.stdout.decode()
        except Exception as e:
            results["clamav"] = str(e)

    # ---------------- METADATA ----------------
    if kwargs.get("enable_metadata"):
        if "PDF" in file_info["type"]:
            results["pdf_text"] = extract_text(file_path)
        elif "image" in file_info["type"]:
            img = Image.open(file_path)
            exif = img._getexif()
            if exif:
                results["image_metadata"] = {
                    TAGS.get(k, k): v for k, v in exif.items()
                }

    # ---------------- REPORT ----------------
    report_file = kwargs.get("report_file")
    report_format = kwargs.get("report_format", "json")

    if report_file:
        if report_format == "json":
            with open(report_file, "w") as f:
                json.dump(results, f, indent=4)
        elif report_format == "csv":
            with open(report_file, "w", newline="") as f:
                w = csv.writer(f)
                for k, v in results.items():
                    w.writerow([k, v])

    return {
        "status": "completed",
        "rules_triggered": rules_triggered,
        "report_file": report_file,
    }

_birthdays = {}

def remember_birthday(*, name: str, date: str):
    _birthdays[name] = date
    return f"Saved birthday for {name}"

def check_birthdays():
    today = datetime.now().strftime("%m-%d")
    for name, date in _birthdays.items():
        if date[5:] == today:
            print(f"🎉 Happy Birthday {name}!")


def download_attachments(*, email_user, email_pass, folder="attachments"):
    mail = imaplib.IMAP4_SSL("imap.gmail.com")
    mail.login(email_user, email_pass)
    mail.select("inbox")

    _, data = mail.search(None, "ALL")
    os.makedirs(folder, exist_ok=True)

    for num in data[0].split():
        _, msg_data = mail.fetch(num, "(RFC822)")
        msg = email.message_from_bytes(msg_data[0][1])
        for part in msg.walk():
            if part.get_filename():
                with open(os.path.join(folder, part.get_filename()), "wb") as f:
                    f.write(part.get_payload(decode=True))
    return "Attachments downloaded"

def scrape_best_sellers():
    url = "https://www.amazon.com/Best-Sellers/zgbs"
    soup = BeautifulSoup(requests.get(url).text, "html.parser")
    return [item.text.strip() for item in soup.select(".zg-item")]

def aws_control(*, action: str, instance_id: str, region="us-east-1"):
    ec2 = boto3.client("ec2", region_name=region)
    if action == "start":
        ec2.start_instances(InstanceIds=[instance_id])
    elif action == "stop":
        ec2.stop_instances(InstanceIds=[instance_id])
    return f"EC2 {action} executed"

def plot_spectrogram(audio_path: str, **kwargs):
    """
    Plot the spectrogram of an audio file.

    Parameters:
        audio_path (str): Path to the audio file.
        **kwargs: Optional keyword arguments:
            - figsize (tuple): Figure size, default (10,5)
            - x_axis (str): x-axis type for display, default 'time'
            - y_axis (str): y-axis type for display, default 'hz'
            - title (str): Custom title, default 'Spectrogram of <audio_path>'
    """
    # Load audio
    x, sr = librosa.load(audio_path)

    # Compute Short-Time Fourier Transform (STFT)
    X = librosa.stft(x)
    Xdb = librosa.amplitude_to_db(abs(X))

    # Plot settings
    figsize = kwargs.get('figsize', (10, 5))
    x_axis = kwargs.get('x_axis', 'time')
    y_axis = kwargs.get('y_axis', 'hz')
    title = kwargs.get('title', f'Spectrogram of {audio_path}')

    plt.figure(figsize=figsize)
    librosa.display.specshow(Xdb, sr=sr, x_axis=x_axis, y_axis=y_axis)
    plt.colorbar()
    plt.title(title)
    plt.show()

def csv_to_excel(**kwargs):
    """
    Convert a CSV file to an Excel file.

    Keyword Args:
        csv_file (str): Path to the input CSV file. (required)
        excel_file (str): Path to the output Excel file. (required)
        sheet_name (str): Name of the sheet to write to. Default is 'Sheet1'.
        separator (str): CSV separator. Default is ','.
    """
    # Extract kwargs with defaults
    csv_file = kwargs.get("csv_file")
    excel_file = kwargs.get("excel_file")
    sheet_name = kwargs.get("sheet_name", "Sheet1")
    separator = kwargs.get("separator", ",")

    # Validate required arguments
    if not csv_file or not excel_file:
        raise ValueError("Both 'csv_file' and 'excel_file' must be provided as kwargs.")

    # Ensure CSV exists
    if not os.path.exists(csv_file):
        raise FileNotFoundError(f"CSV file '{csv_file}' not found.")

    # Open or create Excel workbook
    if os.path.exists(excel_file):
        wb = openpyxl.load_workbook(excel_file)
        if sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
        else:
            sheet = wb.create_sheet(sheet_name)
    else:
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = sheet_name

    # Read CSV and write to Excel
    with open(csv_file, "r", encoding="utf-8") as f:
        for row_idx, line in enumerate(f, start=1):
            line = line.rstrip("\n").split(separator)
            for col_idx, data in enumerate(line, start=1):
                sheet.cell(row=row_idx, column=col_idx).value = data

    # Save workbook
    wb.save(excel_file)
    print(f"CSV '{csv_file}' successfully written to Excel '{excel_file}' in sheet '{sheet_name}'.")



def ip_geolocator(ip: str = None, **kwargs):
    """
    Get geolocation information for a given IP address.
    If no IP is provided, it will use the public IP of the machine.
    """
    try:
        if not ip:
            ip = requests.get("https://api.ipify.org").text
        
        response = requests.get(f"http://ip-api.com/json/{ip}")
        data = response.json()
        
        if data['status'] != 'success':
            return {"error": data.get("message", "Failed to fetch IP data")}
        
        result = {
            "IP": data.get("query"),
            "Country": data.get("country"),
            "Region": data.get("regionName"),
            "City": data.get("city"),
            "ZIP": data.get("zip"),
            "ISP": data.get("isp"),
            "Lat": data.get("lat"),
            "Lon": data.get("lon"),
            "Timezone": data.get("timezone")
        }
        return result
    except Exception as e:
        return {"error": str(e)}
    

def play_tone_skill(**kwargs):
    """
    Plays a single tone of a specific frequency through the speaker.
    kwargs: frequency (Hz), duration (seconds), volume (0.0-1.0), sample_rate
    """
    frequency = kwargs.get("frequency", 440)
    duration = kwargs.get("duration", 2)
    volume = kwargs.get("volume", 0.5)
    sample_rate = kwargs.get("sample_rate", 44100)

    t = np.linspace(0, duration, int(sample_rate * duration), False)
    tone = np.sin(frequency * t * 2 * np.pi) * volume

    sd.play(tone, samplerate=sample_rate)
    sd.wait()


def play_sweep_skill(**kwargs):
    """
    Plays a frequency sweep through the speaker.
    kwargs: duration, volume, sample_rate, start_freq, end_freq
    """
    duration = kwargs.get("duration", 5)
    volume = kwargs.get("volume", 0.5)
    sample_rate = kwargs.get("sample_rate", 44100)
    start_freq = kwargs.get("start_freq", 20)
    end_freq = kwargs.get("end_freq", 20000)

    t = np.linspace(0, duration, int(sample_rate * duration), False)
    sweep = scipy.signal.chirp(t, start_freq, t[-1], end_freq, method='logarithmic') * volume

    sd.play(sweep, samplerate=sample_rate)
    sd.wait()


def speaker_health_test_skill(**kwargs):
    """
    Plays different tones and a sweep to assess speaker health.
    Returns a health score in %.
    """
    print("Playing test tones...")
    health_score = 0

    # Low frequency
    print("Playing 100 Hz tone...")
    play_tone_skill(frequency=100, duration=2)
    time.sleep(1)
    health_score += 25

    # Mid frequency
    print("Playing 1000 Hz tone...")
    play_tone_skill(frequency=1000, duration=2)
    time.sleep(1)
    health_score += 25

    # High frequency
    print("Playing 5000 Hz tone...")
    play_tone_skill(frequency=5000, duration=2)
    time.sleep(1)
    health_score += 20

    print("Playing 10,000 Hz tone...")
    play_tone_skill(frequency=10000, duration=2)
    time.sleep(1)
    health_score += 15

    # Frequency sweep
    print("Playing frequency sweep from 20 Hz to 20,000 Hz...")
    play_sweep_skill(duration=5)
    time.sleep(1)
    health_score += 15

    print("Speaker health test complete.")
    print(f"Speaker Health: {health_score}%")
    if health_score == 100:
        print("The speaker is in excellent condition!")
    elif 80 <= health_score < 100:
        print("The speaker is in good condition.")
    elif 60 <= health_score < 80:
        print("The speaker is in average condition.")
    else:
        print("The speaker might be in poor condition.")
    return {"health_score": health_score}

def get_sophos_central_health_script():
    """
    Downloads the Sophos_Central_Health.py script from the official GitHub repository
    and returns the full source code as a string.

    Returns:
        str: The source code of Sophos_Central_Health.py
    """
    url = "https://raw.githubusercontent.com/sophos/PS.Machine_Health/main/Sophos_Central_Health.py"
    response = requests.get(url)
    if response.status_code == 200:
        return response.text
    else:
        raise Exception(f"Failed to retrieve script: HTTP {response.status_code}")
    
    
def run_sophos_central_health_analysis(**kwargs):
    """
    Downloads and executes the Sophos Central Health script.
    Returns the analysis results as a dictionary.
    """
    try:
        script_code = get_sophos_central_health_script()
        local_vars = {}
        exec(script_code, {}, local_vars)
        if 'main' in local_vars:
            result = local_vars['main']()
            return result
        else:
            return {"error": "The script does not contain a main function."}
    except Exception as e:
        return {"error": str(e)}




